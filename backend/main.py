from __future__ import annotations

import json
import os
import re
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path
from threading import Lock, Thread
from time import sleep
from typing import Any
from uuid import uuid4

from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field, model_validator

try:
    from openai import OpenAI
except Exception:  # pragma: no cover
    OpenAI = None


load_dotenv()

SUPPORTED_EVENT_TYPES = {"wedding", "jubilee"}
PROGRAM_SCHEMA_VERSION = 5
EVENT_AI_MODE = os.getenv("EVENT_AI_MODE", "premium").strip().lower()
AI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.4" if EVENT_AI_MODE == "premium" else "gpt-5.4-mini")
AI_REASONING_EFFORT = os.getenv("OPENAI_REASONING_EFFORT", "medium" if EVENT_AI_MODE == "premium" else "low").strip()
OPENAI_DOSSIER_MODEL = os.getenv("OPENAI_DOSSIER_MODEL", AI_MODEL).strip()
OPENAI_DOSSIER_REASONING = os.getenv("OPENAI_DOSSIER_REASONING", AI_REASONING_EFFORT).strip()
OPENAI_WRITER_MODEL = os.getenv("OPENAI_WRITER_MODEL", "gpt-5.4-mini" if EVENT_AI_MODE == "premium" else AI_MODEL).strip()
OPENAI_WRITER_REASONING = os.getenv("OPENAI_WRITER_REASONING", "medium" if EVENT_AI_MODE == "premium" else AI_REASONING_EFFORT).strip()
STRICT_AI_ONLY = os.getenv("STRICT_AI_ONLY", "true" if EVENT_AI_MODE == "premium" else "false").strip().lower() == "true"
OPENAI_TIMEOUT_SECONDS = float(os.getenv("OPENAI_TIMEOUT_SECONDS", "180" if EVENT_AI_MODE == "premium" else "60"))
OPENAI_TIMEOUT_RETRY_MODEL = os.getenv("OPENAI_TIMEOUT_RETRY_MODEL", "gpt-5.4-mini" if EVENT_AI_MODE == "premium" else "").strip()
OPENAI_TIMEOUT_RETRY_REASONING = os.getenv("OPENAI_TIMEOUT_RETRY_REASONING", "low" if EVENT_AI_MODE == "premium" else AI_REASONING_EFFORT).strip()
OPENAI_POLISH_MODEL = os.getenv("OPENAI_POLISH_MODEL", OPENAI_TIMEOUT_RETRY_MODEL or OPENAI_WRITER_MODEL or AI_MODEL).strip()
OPENAI_POLISH_REASONING = os.getenv("OPENAI_POLISH_REASONING", OPENAI_TIMEOUT_RETRY_REASONING or OPENAI_WRITER_REASONING or AI_REASONING_EFFORT).strip()
OPENAI_GATEWAY_RETRIES = int(os.getenv("OPENAI_GATEWAY_RETRIES", "3"))
OPENAI_GATEWAY_RETRY_DELAY_SECONDS = float(os.getenv("OPENAI_GATEWAY_RETRY_DELAY_SECONDS", "2.5"))
GENERATION_WATCHDOG_SECONDS = float(os.getenv("GENERATION_WATCHDOG_SECONDS", "420"))
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
FRONTEND_ORIGINS = os.getenv("FRONTEND_ORIGINS", "http://localhost:3000")
RAILWAY_VOLUME_MOUNT_PATH = os.getenv("RAILWAY_VOLUME_MOUNT_PATH", "").strip()

DATA_DIR = (
    Path(RAILWAY_VOLUME_MOUNT_PATH) / "event_ai_data"
    if RAILWAY_VOLUME_MOUNT_PATH
    else Path(__file__).resolve().parent / "data"
)
DATA_DIR.mkdir(parents=True, exist_ok=True)
SUBMISSIONS_FILE = DATA_DIR / "submissions.json"
GENERATION_THREAD_LOCK = Lock()
GENERATION_THREADS: dict[str, Thread] = {}

client = OpenAI(api_key=OPENAI_API_KEY, timeout=OPENAI_TIMEOUT_SECONDS, max_retries=1) if OpenAI and OPENAI_API_KEY else None

app = FastAPI(title="Event AI Backend", version="1.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[origin.strip() for origin in FRONTEND_ORIGINS.split(",") if origin.strip()],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class QuestionnaireSubmission(BaseModel):
    eventType: str
    clientName: str = Field(min_length=1)
    phone: str = ""
    eventDate: str = Field(min_length=1)
    city: str = Field(min_length=1)
    venue: str = Field(min_length=1)
    startTime: str = Field(default="")
    guestCount: str = ""
    childrenInfo: str = ""
    atmosphere: str = ""
    fears: str = ""
    hostWishes: str = ""
    references: str = ""
    musicLikes: str = ""
    musicBans: str = ""

    groomName: str = ""
    brideName: str = ""
    weddingTraditions: str = ""
    groomParents: str = ""
    brideParents: str = ""
    grandparents: str = ""
    loveStory: str = ""
    coupleValues: str = ""
    importantDates: str = ""
    proposalStory: str = ""
    nicknames: str = ""
    insideJokes: str = ""
    guestsList: str = ""
    conflictTopics: str = ""
    likedFormats: str = ""
    keyMoments: str = ""

    celebrantName: str = ""
    celebrantAge: str = ""
    familyMembers: str = ""
    anniversaryAtmosphere: str = ""
    biographyStory: str = ""
    achievements: str = ""
    lifeStages: str = ""
    characterTraits: str = ""
    funnyFacts: str = ""
    importantGuests: str = ""
    jubileeConflictTopics: str = ""
    jubileeLikedFormats: str = ""
    whatCannotBeDone: str = ""

    @model_validator(mode="after")
    def validate_business_rules(self) -> "QuestionnaireSubmission":
        if self.eventType not in SUPPORTED_EVENT_TYPES:
            raise ValueError("Поддерживаются только wedding и jubilee")
        if not self.startTime.strip():
            raise ValueError("Поле startTime обязательно")
        if self.eventType == "wedding" and (not self.groomName.strip() or not self.brideName.strip()):
            raise ValueError("Для wedding обязательны groomName и brideName")
        if self.eventType == "jubilee" and not self.celebrantName.strip():
            raise ValueError("Для jubilee обязательно celebrantName")
        return self


def load_submissions() -> list[dict[str, Any]]:
    if not SUBMISSIONS_FILE.exists():
        return []
    try:
        data = json.loads(SUBMISSIONS_FILE.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except (OSError, json.JSONDecodeError):
        return []


def save_submissions(submissions: list[dict[str, Any]]) -> None:
    SUBMISSIONS_FILE.write_text(json.dumps(submissions, ensure_ascii=False, indent=2), encoding="utf-8")


def get_questionnaire_labels() -> dict[str, str]:
    return {
        "eventType": "Тип мероприятия",
        "clientName": "Название заявки",
        "phone": "Телефон",
        "eventDate": "Дата события",
        "city": "Город",
        "venue": "Площадка / ресторан",
        "startTime": "Время начала / сбор гостей",
        "guestCount": "Количество гостей",
        "childrenInfo": "Будут ли дети",
        "atmosphere": "Атмосфера",
        "fears": "Страхи и переживания",
        "hostWishes": "Пожелания к ведущему",
        "references": "Референсы по стилю вечера",
        "musicLikes": "Какая музыка нравится",
        "musicBans": "Что из музыки нельзя включать",
        "groomName": "Имя жениха",
        "brideName": "Имя невесты",
        "weddingTraditions": "Какие свадебные традиции нужны",
        "groomParents": "Как зовут родителей жениха",
        "brideParents": "Как зовут родителей невесты",
        "grandparents": "Бабушки и дедушки",
        "loveStory": "История знакомства",
        "coupleValues": "Главные ценности пары",
        "importantDates": "Важные даты и события пары",
        "proposalStory": "История предложения",
        "nicknames": "Как ласково называют друг друга",
        "insideJokes": "Внутренние шутки / мемы",
        "guestsList": "Имена гостей и характеристики",
        "conflictTopics": "Конфликтные темы / чувствительные фигуры",
        "likedFormats": "Какие форматы нравятся",
        "keyMoments": "Какие 3-5 моментов самые важные",
        "celebrantName": "Имя юбиляра",
        "celebrantAge": "Возраст юбиляра",
        "familyMembers": "Семья юбиляра",
        "anniversaryAtmosphere": "Атмосфера юбилея",
        "biographyStory": "Истории, которые можно использовать",
        "achievements": "Важные достижения",
        "lifeStages": "Важные этапы жизни",
        "characterTraits": "Главные качества юбиляра",
        "funnyFacts": "Любимые фразы / шутки / мемы",
        "importantGuests": "Важные гости",
        "jubileeConflictTopics": "Конфликтные темы / чувствительные фигуры",
        "jubileeLikedFormats": "Какие форматы нравятся",
        "whatCannotBeDone": "Чего нельзя делать",
    }


def build_questionnaire_context(questionnaire: dict[str, Any]) -> str:
    labels = get_questionnaire_labels()
    lines: list[str] = []
    for key, label in labels.items():
        value = str(questionnaire.get(key, "")).strip()
        if value:
            lines.append(f"{label}: {value}")
    return "\n".join(lines)


def list_from_text(value: str, fallback: list[str]) -> list[str]:
    items = [item.strip(" -•\t") for item in re.split(r"[\n,;]+", value or "") if item.strip(" -•\t")]
    return items[:8] if items else fallback


def clean_fragment(value: str, fallback: str = "") -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" .,\n\t")
    return text or fallback


def build_personalization_brief(questionnaire: dict[str, Any]) -> dict[str, Any]:
    is_wedding = questionnaire.get("eventType") == "wedding"
    heroes = (
        f"{questionnaire.get('groomName', '').strip()} и {questionnaire.get('brideName', '').strip()}".strip()
        if is_wedding
        else questionnaire.get("celebrantName", "").strip()
    ) or questionnaire.get("clientName", "").strip() or "герои вечера"
    return {
        "heroes": heroes,
        "love_story": clean_fragment(questionnaire.get("loveStory", ""), "их история не укладывается в банальный пересказ и требует живого человеческого тона"),
        "proposal_story": clean_fragment(questionnaire.get("proposalStory", ""), "в их истории есть момент выбора, который хочется произносить не шаблоном, а с уважением к деталям"),
        "values": list_from_text(questionnaire.get("coupleValues", ""), ["уважение", "чувство юмора", "опора друг на друга"]),
        "nicknames": clean_fragment(questionnaire.get("nicknames", ""), ""),
        "inside_jokes": clean_fragment(questionnaire.get("insideJokes", ""), ""),
        "important_dates": list_from_text(questionnaire.get("importantDates", ""), []),
        "key_moments": list_from_text(questionnaire.get("keyMoments", ""), ["сильное открытие", "эмоциональное ядро", "светлый финал"]),
        "host_wishes": clean_fragment(questionnaire.get("hostWishes", ""), ""),
        "references": clean_fragment(questionnaire.get("references", ""), ""),
        "family": {
            "groom_parents": clean_fragment(questionnaire.get("groomParents", ""), ""),
            "bride_parents": clean_fragment(questionnaire.get("brideParents", ""), ""),
            "grandparents": clean_fragment(questionnaire.get("grandparents", ""), ""),
            "family_members": clean_fragment(questionnaire.get("familyMembers", ""), ""),
        },
        "guests_texture": clean_fragment(
            questionnaire.get("guestsList", "") or questionnaire.get("importantGuests", ""),
            "в зале есть люди, для которых эта история действительно личная, а не формальная",
        ),
        "music_likes": extract_music_preferences(questionnaire),
    }


def parse_event_date_context(questionnaire: dict[str, Any]) -> dict[str, Any]:
    raw_date = str(questionnaire.get("eventDate", "")).strip()
    month = ""
    season = "all-season"
    year = datetime.now().year
    for pattern in ("%Y-%m-%d", "%d.%m.%Y", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            parsed = datetime.strptime(raw_date, pattern)
            year = parsed.year
            month = parsed.strftime("%B")
            if parsed.month in {12, 1, 2}:
                season = "winter"
            elif parsed.month in {3, 4, 5}:
                season = "spring"
            elif parsed.month in {6, 7, 8}:
                season = "summer"
            else:
                season = "autumn"
            break
        except ValueError:
            continue
    return {
        "raw_date": raw_date,
        "event_year": year,
        "event_month": month,
        "event_season": season,
    }


def build_trend_bank(questionnaire: dict[str, Any]) -> dict[str, Any]:
    event_context = parse_event_date_context(questionnaire)
    event_year = event_context["event_year"]
    is_wedding = questionnaire.get("eventType") == "wedding"
    season = event_context["event_season"]

    seasonal_logic = {
        "winter": "больше кинематографичности, теплого света, камерности и контрастной музыкальной драматургии",
        "spring": "больше ощущения новизны, воздуха, editorial-lightness и романтики без сиропа",
        "summer": "больше живого движения, легкости, open-air energy, позднего танцпола и естественного контакта с гостями",
        "autumn": "больше глубины, фактуры, благородного тепла, взрослой эмоции и насыщенного narrative-ритма",
        "all-season": "логика подбирается по анкете, но с приоритетом персонализации и режиссерской цельности",
    }

    wedding_trends = [
        "editorial storytelling вместо банкетного набора тостов и конкурсов",
        "гостевой experience строится через curated moments, а не через постоянное микрофонное давление",
        "audio guest book / voice notes / короткие аудио-включения от близких работают сильнее, чем длинная череда одинаковых поздравлений",
        "portrait-studio и documentary-подход к гостям усиливают ощущение события как красивой личной истории, а не просто банкета",
        "диджей работает сетами и переходами по драматургии, а не жанровыми общими словами",
        "любимые треки пары ставятся как sacred points вечера, а не как случайные пожелания в конце списка",
    ]
    jubilee_trends = [
        "биографическая режиссура сильнее работает через короткие точные истории, чем через линейный пересказ жизни",
        "уважение к статусу гостей и героя вечера строится через вкус и точность, а не через официоз",
        "музыкальная логика юбилея должна держать баланс между узнаваемостью и современным качеством подачи",
        "вовлечение строится через адресность, семейный контекст и наблюдательность, а не через сценическое давление",
    ]

    sources = [
        "The Knot, wedding technology and guest-experience trends for 2026 planning logic, проверено 2026-04-15",
        "Brides, portrait studio and editorial guest-experience trend, проверено 2026-04-15",
        "Los Angeles Times, wedding entertainment trends and guest-engagement direction, проверено 2026-04-15",
        "THE WED, wedding content creation and documentary-style storytelling trends, проверено 2026-04-15",
    ]

    return {
        "mode": EVENT_AI_MODE,
        "event_year": event_year,
        "season": season,
        "seasonal_logic": seasonal_logic.get(season, seasonal_logic["all-season"]),
        "applied_market_direction": wedding_trends if is_wedding else jubilee_trends,
        "dj_market_direction": [
            "welcome и dinner строятся через curated tone-setting, а не через просто 'спокойную музыку'",
            "танцпол поднимается волнами: familiar lift -> stylish singalong -> peak -> elegant close",
            "переход из трогательного блока должен быть ступенчатым, без ломаного эмоционального удара",
            "сеты должны учитывать возраст, город, площадку и любимые треки пары как драматургические маркеры",
        ],
        "host_market_direction": [
            "ведущий работает как storyteller-editor, а не как человек, объявляющий блоки",
            "реплики должны звучать как авторский текст под конкретных людей, а не как красивая общая заготовка",
            "мягкое вовлечение важнее шумной активности: наблюдательность, точность, персональный юмор и уместность",
        ],
        "sources_checked_at": "2026-04-15",
        "sources": sources,
    }


def build_style_bank(questionnaire: dict[str, Any]) -> dict[str, Any]:
    is_wedding = questionnaire.get("eventType") == "wedding"
    profile = build_personalization_brief(questionnaire)
    event_context = parse_event_date_context(questionnaire)
    return {
        "mode": EVENT_AI_MODE,
        "voice_direction": (
            "expensive editorial romance with wit, nerve and real humanity"
            if is_wedding
            else "warm prestige with wit, biography depth and living stage presence"
        ),
        "must_sound_like": [
            "авторская сценическая речь, написанная под конкретных людей",
            "живой режиссерский текст, а не обзор анкеты",
            "мягкая интеллектуальная сила без пафоса и канцелярита",
        ],
        "must_not_sound_like": [
            "шаблонная свадебная открытка",
            "банкетный универсальный текст",
            "пересказ полей анкеты почти теми же словами",
            "общие слова про атмосферность, любовь и судьбу без конкретной фактуры",
        ],
        "negative_prompt_lexicon": [
            "данное мероприятие",
            "виновники торжества",
            "взаимное уважение",
            "на сегодняшний день",
            "наш профессиональный ведущий",
            "согласно анкете",
            "мы проанализировали",
            "мероприятие",
            "идеальная пара",
            "две половинки",
        ],
        "metaphor_rules": [
            "метафоры должны вытекать из фактуры пары, а не жить отдельно от нее",
            "если образ не помогает сцене и не усиливает конкретных героев, его не использовать",
            "в каждом длинном тексте должна быть не только красота, но и мысль, темперамент и режиссерский контроль",
            "использовать одну центральную метафору и 2-3 вспомогательных образа, а не бесконечно повторять один и тот же образ",
        ],
        "personalization_rules": [
            f"обязательно встроить реальные детали пары: {profile['love_story']}",
            f"обязательно встроить ценности и внутренний язык пары: {', '.join(profile['values'][:3])}; {profile['nicknames'] or 'личные обращения, если они уместны'}",
            f"если есть внутренние шутки, использовать их точечно и умно: {profile['inside_jokes'] or 'не выдумывать шутки без опоры на анкету'}",
            "минимум четыре блока должны содержать детали, которые нельзя перенести в чужую анкету без явной фальши",
        ],
        "tone_roles": [
            "стендап-комик, который верит в любовь",
            "рок-поэт, создающий гимн для двоих",
            "сценарист-драматург с чувством ритма, образа и уместной иронии",
        ],
        "detail_to_meaning_formula": [
            "факт из анкеты -> смысл -> сценическая подача",
            "деталь пары -> метафора -> эмоциональный призыв",
            "личная фактура -> режиссерский вывод -> готовая речь ведущего",
        ],
        "dj_rules": [
            "лист диджея строить как реальный sheet по точкам вечера, а не как набор жанров",
            "обязательные треки и любимые песни клиента указывать в конкретных слотах и с задачей блока",
            "для каждого большого музыкального участка задавать energy arc, переход и stop list",
        ],
        "date_awareness": f"событие запланировано на {event_context['raw_date'] or 'указанную дату'}; использовать рыночную логику именно для {event_context['event_year']} года и сезона {event_context['event_season']}",
    }


def infer_central_metaphor(questionnaire: dict[str, Any]) -> dict[str, Any]:
    profile = build_personalization_brief(questionnaire)
    combined = " ".join(
        [
            profile["love_story"].lower(),
            profile["proposal_story"].lower(),
            profile["references"].lower(),
            " ".join(profile["music_likes"]).lower(),
            str(questionnaire.get("loveStory", "")).lower(),
            str(questionnaire.get("proposalStory", "")).lower(),
            str(questionnaire.get("references", "")).lower(),
            str(questionnaire.get("musicLikes", "")).lower(),
        ]
    )
    if any(word in combined for word in ["лес", "гриб", "тропа", "чаща"]):
        return {
            "central_metaphor": "любовь как тихая охота за сокровищем в лесу",
            "image_dictionary": [
                "заповедные тропы судьбы",
                "тихая охота за счастьем",
                "лесная глубина, в которой находят самое ценное",
                "корзина редких находок вместо случайных совпадений",
            ],
        }
    if any(word in combined for word in ["2015", "десять лет", "год", "вместе с"]):
        return {
            "central_metaphor": "отношения как эпоха, выдержанная временем",
            "image_dictionary": [
                "выдержанное вино",
                "механизм, который годами работает точнее",
                "эпоха двоих",
                "долгая пластинка, которая не теряет глубины звучания",
            ],
        }
    if any(word in combined for word in ["музы", "музык", "рок", "ленинград", "нога свело", "концерт"]):
        return {
            "central_metaphor": "любовь как живая песня с характером, нервом и своим ритмом",
            "image_dictionary": [
                "риф, который узнается с первых секунд",
                "припев, в который хочется вернуться",
                "громкость сердца без фальшивых нот",
                "общий ритм, который не собьешь случайным шумом",
            ],
        }
    return {
        "central_metaphor": "любовь как авторская история, которая собирается из редких деталей и сильных решений",
        "image_dictionary": [
            "роман с собственным почерком",
            "кадр, который нельзя спутать с чужим",
            "свет, найденный не на витрине, а внутри",
            "история, сшитая по фигуре этих двоих",
        ],
    }


def build_dramaturgy_bank(questionnaire: dict[str, Any]) -> dict[str, Any]:
    profile = build_personalization_brief(questionnaire)
    metaphor = infer_central_metaphor(questionnaire)
    return {
        "role_shift": "агент работает не как аналитик ТЗ, а как сценарист-драматург и креативный помощник",
        "central_metaphor": metaphor["central_metaphor"],
        "image_dictionary": metaphor["image_dictionary"],
        "unique_hooks": [
            profile["love_story"],
            profile["proposal_story"],
            profile["inside_jokes"] or "найти одну живую, небанальную интонацию пары",
            ", ".join(profile["important_dates"][:3]) or "вшить реальные точки их общего времени",
            ", ".join(profile["music_likes"][:3]) or "использовать музыкальные вкусы как часть характера пары",
        ],
        "opening_variants": [
            "эмоциональный",
            "драйвовый",
            "метафоричный",
        ],
        "writing_formula": [
            "деталь -> смысл -> сценическая подача",
            "факт пары -> образ -> живая реплика в зал",
            "личная зацепка -> драматургический вывод -> готовый монолог ведущего",
        ],
    }


def add_minutes(time_value: str, minutes: int) -> str:
    try:
        base_time = datetime.strptime(time_value.strip(), "%H:%M")
    except ValueError:
        base_time = datetime.strptime("18:00", "%H:%M")
    anchored = base_time.replace(year=2000, month=1, day=1)
    return (anchored + timedelta(minutes=minutes)).strftime("%H:%M")


def build_metaphoric_text(
    *,
    event_label: str,
    heroes: str,
    atmosphere: str,
    block_label: str,
    purpose: str,
    guest_focus: str,
) -> str:
    metaphors = [
        f"{block_label} должен войти в зал как первый луч в распахнутое окно и сразу показать людям, куда сегодня смотрит сердце вечера.",
        "Слова ведущего здесь работают не как кирпичи протокола, а как мягкий свет на плечах гостей: они не давят, а собирают внимание.",
        "Настроение важно собирать как оркестр перед увертюрой: сначала дыхание, потом ритм, потом общая мелодия зала.",
        "Этот момент должен стать не формальной остановкой, а мостом между людьми, по которому гости спокойно переходят к следующей эмоции.",
        f"История {heroes} должна звучать не как отполированная открытка, а как река с течением, глубиной и настоящими берегами.",
        f"Атмосфера {atmosphere} здесь нужна как дорогой аромат: ее невозможно взять руками, но именно она делает вечер узнаваемым.",
        "Каждый поворот речи пусть работает как камера крупного плана, вытаскивая не шум, а ту самую деталь, из-за которой вечер становится личным.",
        "Зал в этом блоке важно держать как единый организм: без суеты, без лишнего шума, без ощущения, что людей тянут за локоть в чужую эмоцию.",
        "Ведущий держит ритм как дирижер держит паузу: уверенно, спокойно и с пониманием, что тишина иногда звучит громче фанфар.",
        "Финальная интонация блока должна закрываться как точная кинематографичная склейка, после которой следующий эпизод начинается естественно.",
        "Хорошая реплика здесь должна сработать как бокал тонкого стекла: легкая на слух, но с долгим послевкусием.",
        "Герои вечера не витрина, а живой нерв события, и задача речи не полировать его до глянца, а подсветить так, чтобы зал это почувствовал.",
    ]
    return " ".join(
        [
            f"Сейчас задача ведущего не просто объявить {block_label.lower()}, а сделать настоящую сценическую работу: {purpose}.",
            f"Фокус блока — {guest_focus}.",
            *metaphors,
            f"Именно так {event_label.lower()} перестает быть набором пунктов и превращается в живую рабочую драматургию для реальной площадки.",
        ]
    )


def build_detailed_timeline(
    questionnaire: dict[str, Any], heroes: str, atmosphere: str, hard_bans: list[str]
) -> list[dict[str, str]]:
    start = questionnaire.get("startTime", "").strip() or "18:00"
    event_label = "свадьба" if questionnaire["eventType"] == "wedding" else "юбилей"
    blocks = [
        ("Сбор гостей и тонкий welcome", 20, "собрать внимание без суеты и создать ощущение дорогого, бережного старта", "снять стартовое напряжение и дать гостям почувствовать уверенность команды", "мягкое приветствие, ориентирование гостей, короткие точные касания", "держать стильный фон без вокального перегруза и резких скачков", "не перегружать микрофон в первые минуты", "когда зал физически собран, перевести внимание в официальное открытие"),
        ("Главное открытие вечера", 15, "задать большую идею события и закрепить эмоциональный код вечера", "вывести в центр истории героев и быстро создать единое поле внимания", "короткий сбор тишины, сильная открывающая речь, ясная рамка вечера", "подвести кинематографичный старт и аккуратно убрать музыку под голос", "не растянуть открытие дольше, чем держит внимание зала", "после аплодисментов открыть дорогу к первому смысловому ядру"),
        ("Первое смысловое ядро", 20, "дать залу почувствовать живую персонализацию и не уйти в шаблон", "закрепить, что этот вечер построен вокруг реальных людей, а не формата", "подводка к истории пары или героя вечера, включение первых личных смыслов", "держать теплый фон и не перебивать личные детали клишированными подводками", "не превращать историю в пересказ анкеты", "через теплый акцент перевести внимание к семье и близким"),
        ("Семейный и эмоциональный блок", 25, "поставить в центр близких людей и собрать первое глубокое чувство зала", "дать родным безопасное пространство для сильных слов", "подводка к родителям, близким, семейным смыслам и благодарности", "держать благородную эмоциональную музыку без сахарной банальности", "не торопить чувства и не давить на чувствительные темы", "после кульминации мягко поднять зал из глубины в свет"),
        ("Мягкое вовлечение гостей", 20, "оживить зал без кринжа и без давления на интровертов", "включить активных гостей и оставить чувство уважения у скромных", "короткий стильный интерактив, наблюдение, адресные реплики, мягкая динамика", "дать ритм средней энергии и убрать все, что похоже на устаревшие конкурсы", f"не допускать {', '.join(hard_bans[:3])}", "после оживления подготовить точку для сюрприза или специального выхода"),
        ("Сюрприз или специальный акцент", 20, "создать эффект неожиданности, который усиливает историю вечера, а не ломает ее", "дать гостям ощущение эксклюзивности и собранности драматургии", "подводка к сюрпризу, номеру, видео или специальному поздравлению", "сделать чистый музыкальный коридор на вход и выход сюрприза", "не раскрыть интригу раньше времени и не потерять темп после нее", "после акцента аккуратно перевести зал в танцевальную волну"),
        ("Танцевальная волна", 35, "поднять энергию и дать залу физически прожить вечер", "включить людей в движение без ощущения принуждения", "подводка к первому сильному танцевальному выходу и управляемый разгон", "разгонять пол постепенно, треками-магнитами, а не резким ударом в потолок", "не выжечь энергию слишком рано", "на пике собрать внимание и подготовить финальный блок"),
        ("Финал и закрытие вечера", 15, "собрать вечер в ясную красивую точку и оставить послевкусие", "закрыть эмоцию, поблагодарить, дать людям ощущение завершенности", "финальная речь, общий акцент, благодарность и смысловой последний кадр", "дать объединяющий финальный трек без хаотичного добивания", "не растягивать финал после смысловой точки", "закрыть вечер так, чтобы гости унесли его внутри, а не только на фото"),
    ]

    timeline: list[dict[str, str]] = []
    elapsed = 0
    for title, duration, purpose, focus, action, dj_task, risk, transition in blocks:
        time_from = add_minutes(start, elapsed)
        elapsed += duration
        time_to = add_minutes(start, elapsed)
        timeline.append(
            {
                "time_from": time_from,
                "time_to": time_to,
                "block_title": title,
                "block_purpose": purpose,
                "what_happens": action,
                "host_action": focus,
                "host_text": build_metaphoric_text(
                    event_label=event_label,
                    heroes=heroes,
                    atmosphere=atmosphere,
                    block_label=title,
                    purpose=purpose,
                    guest_focus=focus,
                ),
                "dj_task": dj_task,
                "director_move": f"Режиссерский ход блока: {purpose}. Точка управления вниманием: {focus}.",
                "risk_control": risk,
                "transition": transition,
            }
        )
    return timeline


def build_detailed_host_script(questionnaire: dict[str, Any], heroes: str, atmosphere: str) -> dict[str, str]:
    event_label = "свадьба" if questionnaire["eventType"] == "wedding" else "юбилей"
    return {
        "opening_main": build_metaphoric_text(event_label=event_label, heroes=heroes, atmosphere=atmosphere, block_label="главное открытие", purpose="открыть вечер сильной, образной и пригодной к реальному чтению речью", guest_focus="создать единое дыхание зала и быстро собрать уважительное внимание"),
        "opening_short": build_metaphoric_text(event_label=event_label, heroes=heroes, atmosphere=atmosphere, block_label="короткое открытие", purpose="дать короткую версию открытия без потери образности и статуса", guest_focus="собрать внимание быстро, если тайминг уже напряжен"),
        "welcome_line": build_metaphoric_text(event_label=event_label, heroes=heroes, atmosphere=atmosphere, block_label="welcome-линия", purpose="мягко встретить зал и обозначить интонацию вечера с первых секунд", guest_focus="дать гостям ощущение, что команда управляет атмосферой бережно"),
        "first_core_intro": build_metaphoric_text(event_label=event_label, heroes=heroes, atmosphere=atmosphere, block_label="первое смысловое ядро", purpose="подвести зал к личной истории и первой глубокой точке вечера", guest_focus="перевести внимание от формы к настоящим людям"),
        "family_block_intro": build_metaphoric_text(event_label=event_label, heroes=heroes, atmosphere=atmosphere, block_label="семейный блок", purpose="открыть пространство для близких людей без сладкой банальности", guest_focus="дать родным силу и безопасность для искренних слов"),
        "surprise_intro": build_metaphoric_text(event_label=event_label, heroes=heroes, atmosphere=atmosphere, block_label="сюрприз", purpose="ввести неожиданный поворот красиво и без разрушения общей драматургии", guest_focus="сохранить интригу и удержать зал в доверии"),
        "dance_block_intro": build_metaphoric_text(event_label=event_label, heroes=heroes, atmosphere=atmosphere, block_label="танцевальный блок", purpose="поднять энергию зала современно и без ощущения деревенского форсажа", guest_focus="дать людям разрешение включиться свободно и с удовольствием"),
        "final_block_intro": build_metaphoric_text(event_label=event_label, heroes=heroes, atmosphere=atmosphere, block_label="финальный блок", purpose="собрать все смыслы вечера перед закрытием", guest_focus="оставить у гостей чувство завершенности и ценности пережитого"),
        "closing_words": build_metaphoric_text(event_label=event_label, heroes=heroes, atmosphere=atmosphere, block_label="закрытие вечера", purpose="закрыть событие длинной, сценичной, финальной речью, которую можно реально читать со сцены", guest_focus="оставить теплое послевкусие и внутреннюю точку у каждого гостя"),
    }


def build_direct_host_texts(questionnaire: dict[str, Any], heroes: str, atmosphere: str) -> dict[str, str]:
    event_word = "свадьба" if questionnaire["eventType"] == "wedding" else "юбилей"
    city = questionnaire.get("city", "").strip()
    venue = questionnaire.get("venue", "").strip()
    profile = build_personalization_brief(questionnaire)
    values = ", ".join(profile["values"][:3])
    key_moments = ", ".join(profile["key_moments"][:3])
    important_dates = ", ".join(profile["important_dates"][:3]) or "их собственные важные точки пути"
    nicknames = profile["nicknames"] or "их личный язык двоих"
    inside_jokes = profile["inside_jokes"] or "внутренние смешные коды, понятные самым близким"
    family_line = ", ".join(
        [
            item
            for item in [
                profile["family"]["groom_parents"],
                profile["family"]["bride_parents"],
                profile["family"]["grandparents"],
                profile["family"]["family_members"],
            ]
            if item
        ]
    ) or "их семья и самые близкие люди"
    host_wishes = profile["host_wishes"] or "вести вечер тонко, живо и без банкетной фальши"
    references = profile["references"] or "ориентир на современный, вкусный и взрослый тон"
    love_story = profile["love_story"]
    proposal_story = profile["proposal_story"]
    music_anchor = ", ".join(profile["music_likes"][:3]) or "их любимая музыка и личные ассоциации"
    return {
        "Сбор гостей и тонкий welcome": (
            f"Добрый вечер, дорогие друзья. Сегодня мы начинаем не с формальности, а с правильного воздуха. Пусть этот зал {venue or 'сегодняшней площадки'} сначала наполнится не шумом, а настроением, которое подходит именно истории {heroes}. "
            f"Потому что у этой пары есть не абстрактная романтическая рамка, а свой собственный почерк: {love_story}. И с первых минут важно почувствовать, что мы входим не в стандартный банкетный вечер, а в пространство, где слышны их ценности: {values}. "
            f"Пусть каждый, кто входит сюда, чувствует не банкетную суету, а ощущение, будто вечер уже держит его за руку и спокойно ведет внутрь истории. Пусть первые минуты будут как мягкий свет на стекле, как тихая настройка оркестра перед увертюрой, как первый вдох перед красивой фразой, как шелковая лента на ветру, как теплый янтарь в бокале, как легкий ток воздуха перед открытым окном, как спокойная вода перед отражением огней, как уверенный шаг по сцене без лишнего жеста, как улыбка без нажима, как музыка, которая не вторгается, а собирает людей в единый рисунок. "
            f"Мы никуда не спешим, потому что хорошая {event_word} начинается не с громкости, а с точности. С точности к их личному языку, в котором есть {nicknames}. С точности к тону, которого они ждут от нас: {host_wishes}. С точности к атмосфере {atmosphere}, которая должна ощущаться не на уровне декора, а на уровне внутреннего движения зала. "
            f"И пока зал собирается, давайте позволим этому вечеру стать не толпой голосов, а живым пространством, где каждая деталь уже работает на историю именно этих двоих, а не на очередную красивую картинку без имени."
        ),
        "Главное открытие вечера": (
            f"Добрый вечер, друзья. Сегодня здесь собрались не просто гости, а свидетели конкретной, упрямой, живой истории {heroes}. И если попытаться уложить ее в одну общую открытку, мы сразу потеряем главное. Потому что у этой пары есть свои даты, свои коды, свои опорные слова, свои решения и свои повороты: {important_dates}. "
            f"Есть вечера, которые похожи на список обязательных действий, а есть вечера, которые похожи на хорошо написанную музыку: в них есть вступление, дыхание, кульминация и тот самый аккорд, после которого люди молчат не потому, что не знают, что сказать, а потому, что чувствуют больше, чем можно уместить в одну реплику. Пусть этот вечер будет именно таким. "
            f"Не открыткой под стеклом, а рекой с течением. Не набором тостов, а световым коридором, по которому мы вместе пройдем к самым важным словам. Не салютом ради салюта, а точной режиссурой сердца. Не картонной романтикой, а живым током между людьми. Не чужим шаблоном, а пиджаком, сшитым точно по фигуре этой истории. Не банкетной машиной, а сценой, где каждая деталь знает свое место. Не громкой витриной, а глубоким кадром. Не пестрым шумом, а дорогим ритмом. Не формальным вечером, а настоящим событием. "
            f"Поэтому сегодня мы будем беречь важные паузы, ценить точные слова, смеяться там, где легко, и быть честными там, где важно. И делать это в тоне, который выбрали сами {heroes}: {references}. "
            f"И если к финалу у каждого внутри останется не общее впечатление от красивой программы, а ясное ощущение, что он побывал внутри истории именно этих двоих, значит все было сделано правильно."
        ),
        "Первое смысловое ядро": (
            f"Давайте сейчас не торопиться дальше по программе, а на минуту приблизим фокус. Потому что {heroes} — это не дата в календаре и не парадный кадр для фотографа. Это история, у которой есть собственный почерк, собственный темп и та особенная интонация, которую невозможно придумать в общем шаблоне. "
            f"Если прислушаться к этой истории внимательно, в ней уже слышно, что именно делает ее узнаваемой: {love_story}. А еще в ней есть момент, который особенно важно произносить не формально, а точно: {proposal_story}. "
            f"Иногда любовь похожа на тихую гавань, иногда на дерзкий ветер в лицо, иногда на настольную лампу, под которой в два часа ночи все наконец становится честным, иногда на мост, который держится не на бетоне, а на доверии, иногда на старый винил с теплым потрескиванием, иногда на город после дождя, когда асфальт вдруг начинает отражать свет, иногда на смех, который снимает тяжесть лучше длинных речей, иногда на компас, который не кричит, а просто показывает направление, иногда на сад, который требует ухода, иногда на дом, где даже молчание звучит как свой язык. "
            f"Именно поэтому сегодня важно не рассказывать про вас общими словами, а почувствовать, что у этой истории есть ткань, температура и характер. В ней есть {inside_jokes}. В ней есть важные точки, которые знают самые близкие: {important_dates}. В ней есть то, что вы особенно хотите сохранить в этом вечере: {key_moments}. "
            f"Пусть зал услышит не пересказ анкеты, а настоящее биение этой истории. Пусть каждый сейчас поймет: здесь празднуют не идеальную картинку, а живую связь, которая выдерживает время, быт, расстояния, усталость, смех и все то, из чего и строится настоящая близость."
        ),
        "Семейный и эмоциональный блок": (
            "Сейчас наступает момент, к которому нельзя подходить с пластмассовыми словами. "
            f"Потому что рядом с историей {heroes} всегда стояли люди, без которых этот вечер не звучал бы так полно. И сегодня это не абстрактное слово «семья», а конкретные люди: {family_line}. "
            "Есть близкие, которые похожи на корни большого дерева: их не всегда видно с первого взгляда, но именно они удерживают силу в бурю. "
            "Есть люди, похожие на тихий свет в окне, к которому всегда возвращаются. Есть люди, похожие на камертон, выравнивающий внутренний звук. Есть люди, похожие на старый мост, по которому можно пройти даже в самую трудную погоду. Есть люди, похожие на книгу с закладками на нужных страницах. Есть люди, похожие на мягкий плед в слишком резкий день. Есть люди, похожие на якорь, который не мешает движению, а защищает от дрейфа. Есть люди, похожие на семейное серебро: не кричат о своей ценности, но делают дом домом. Есть люди, похожие на садовую землю: в ней не видно будущих цветов, но без нее не вырастет ничего. Есть люди, похожие на дыхание за кулисами перед выходом: их не всегда замечают, но без них сцена не держится. "
            f"И сегодня самое важное — не сыграть чувство, а дать ему место. Дать место памяти о том, как формировались {values}, как рядом с этой парой стояли близкие, как именно они помогали истории {heroes} становиться прочнее. "
            "Пусть слова, которые сейчас прозвучат, будут не дежурными, а настоящими. "
            "Пусть в них будет благодарность без сиропа, любовь без штампа и тепло без лишней театральности. "
            "Потому что самые сильные семейные слова никогда не шумят. Они просто остаются внутри надолго."
        ),
        "Мягкое вовлечение гостей": (
            "А теперь давайте сделаем то, что отличает живой вечер от красивой, но неподвижной декорации. "
            "Давайте впустим в него зал. Но не шумно, не грубо, не по-банкетному в лоб, а легко, точно и с уважением. "
            "Хорошее вовлечение — это не когда человека вытянули в центр и заставили изображать энтузиазм. Хорошее вовлечение — это когда человек сам улыбается, потому что почувствовал: его не используют, его увидели. "
            "Пусть этот момент будет как искра в камине, а не как петарда в коридоре. Как удачный импровизационный джаз, а не как строевая команда. Как дверь, которая открывается мягко, а не как сирена. Как бокал шампанского, а не как пластиковый мегафон. Как светский разговор с юмором, а не как допрос под софитом. Как шаг на танцпол по собственному желанию, а не как принудительная эвакуация веселья. Как дружеское подмигивание, а не как атака. Как тонкая режиссура, а не как сельский форсаж. Как точный штрих, а не как пятно. Как живая волна, а не как неловкая команда “всем срочно веселиться”. "
            "Поэтому сейчас мы включаем зал бережно: тех, кто готов, тех, кто улыбается, тех, кто умеет поддержать атмосферу. "
            "И делаем это так, чтобы всем остальным захотелось присоединиться добровольно."
        ),
        "Сюрприз или специальный акцент": (
            "У хорошего вечера всегда есть момент, который открывается как потайная дверь в стене знакомой комнаты. "
            "Кажется, ты уже понимаешь пространство, но вдруг оно становится глубже, интереснее и богаче. "
            "Сейчас как раз такой момент. "
            "Пусть он войдет не как случайная вставка в тайминг, а как точный режиссерский поворот. Как письмо, найденное в кармане старого пальто. Как дополнительная комната в красивом доме. Как неожиданный свет в конце коридора. Как виниловая пластинка с редкой версией любимой песни. Как ночной город, который вдруг открывается с новой высоты. Как тонкая реплика в хорошем кино, после которой весь предыдущий сюжет начинает звучать глубже. Как подарок без крика, но с весом. Как драгоценная пуговица на идеально скроенном пиджаке. Как мост между тем, что уже прожито сегодня, и тем, что еще только раскроется. "
            f"Сейчас не нужно объяснять слишком много. Нужно открыть этот момент чисто, красиво и вовремя, опираясь на то, что действительно важно для {heroes}: {key_moments}. "
            "И когда он прозвучит, давайте не торопиться его сбить. "
            "Хороший сюрприз не требует шума. Он требует точной паузы до и уважительной тишины после."
        ),
        "Танцевальная волна": (
            "А теперь пришло время, когда вечер должен перестать быть только историей на словах и стать историей в теле. "
            "Потому что хороший праздник однажды обязательно выходит из текста в движение. "
            "Не резко, не криком, не командой “поднялись все”, а так, будто музыка сама открывает в полу скрытую пружину. "
            "Пусть танцпол сейчас сработает как магнитное поле, как разогревшийся винил, как городская площадь в теплую ночь, как электричество в летнем воздухе перед грозой, как раскрытый парус, который наконец поймал ветер, как неон на мокром асфальте, как мотор, который мурлыкал вхолостую и вдруг пошел ровно, как лестница без ступенчатого страха, как пульс, который из спокойного стал праздничным, как волна, в которую входят не из-под палки, а потому что она зовет. "
            f"Поэтому сейчас мы не ломаем людей в танец, а приглашаем их в энергию. И приглашаем через музыку, которая имеет для этой пары личный вес: {music_anchor}. "
            "Сначала тех, кто готов быть первыми и смелыми, потом тех, кому нужен один знакомый трек, чтобы разрешить себе шаг вперед. "
            "И если все сделано точно, то через несколько минут зал уже не уговаривают. Он дышит музыкой сам."
        ),
        "Финал и закрытие вечера": (
            f"У каждого хорошего вечера есть момент, когда становится ясно: дальше ничего не нужно добавлять, иначе можно испортить рисунок. "
            f"И сейчас мы подходим именно к такой точке. "
            f"Сегодняшняя история {heroes} уже успела пройти через свет, смех, тепло, глубину, музыку, движение и те слова, которые не всегда легко произносить вслух. "
            "Теперь важно закрыть ее красиво. Не как выключают свет в зале после аренды, а как закрывают книгу, к которой еще захочется вернуться. Как ставят последнюю точку в письме, которое будут хранить. Как опускают занавес после спектакля, в котором не было лишних сцен. Как оставляют последний мазок на холсте. Как допивают редкое вино не залпом, а с пониманием вкуса. Как отпускают в ночь лодку с фонарем. Как собирают со стола не остатки, а послевкусие. Как снимают иглу с пластинки в нужную секунду. Как уходят со сцены не потому, что кончилось время, а потому что сказано главное. Как смотрят на город с высоты и понимают: этот кадр запомнится. "
            f"Поэтому сейчас хочется сказать просто и точно: спасибо за этот вечер, за ваше внимание, за ваши сердца, за ваш смех, за вашу деликатность в нужные моменты и за ту живую энергию, которая сделала этот день настоящим событием вокруг истории {heroes}. "
            f"Пусть все самое важное, что сегодня прозвучало, останется не только в фотографиях и видео, но и внутри: их ценности {values}, их собственный язык {nicknames}, их хрупкие и сильные моменты, их живая человеческая правда. "
            "И если завтра вы вспомните не только детали программы, а само ощущение красоты и смысла, значит этот вечер прожит правильно."
        ),
    }


def build_precise_dj_guidance(questionnaire: dict[str, Any], hard_bans: list[str]) -> dict[str, Any]:
    is_wedding = questionnaire["eventType"] == "wedding"
    return {
        "overall_music_policy": (
            "Музыкальная драматургия строится волнами. Welcome держим в зоне 98-112 BPM, открытие — атмосферный спад под голос, "
            "эмоциональные сцены — без жирного вокала и навязчивого бита, первый заход на танцпол — через узнаваемые mid-energy хиты, "
            "главный разгон — только после того, как зал физически вышел и принял общий ритм."
        ),
        "welcome_music": (
            "Ориентиры: Roosevelt `Feels Right`, Purple Disco Machine `Hypnotized`, Parcels `Somethinggreater`, "
            "Jungle `Back On 74`, SG Lewis `Warm`, Satin Jackets `You Make Me Feel Good`. "
            "Задача: создать ощущение дорогого, современного, легкого старта без банкетного шума."
        ),
        "opening_music": (
            "Подводка на открытие: instrumental / cinematic electronic. Ориентиры: RY X `Berlin` (instrumental mood), "
            "Ólafur Arnalds piano textures, ambient intro в духе The xx / London Grammar без плотного вокала. "
            "За 8-12 секунд до первых слов ведущего аккуратно убрать низ и оставить чистый коридор под голос."
        ),
        "table_background": (
            "Фон на столы: Sade, Rhye, Norah Jones, Michael Kiwanuka, Jessie Ware, Leon Bridges, FKJ. "
            "Не играть радио-хиты лоб в лоб. Музыка должна поддерживать разговоры, а не перетягивать на себя внимание."
        ),
        "emotional_blocks_music": (
            "Эмоциональные сцены: Ludovico Einaudi, Ólafur Arnalds, Max Richter, Dustin O'Halloran, ambient piano, soft strings. "
            "Выход из трогательного блока делать через теплый mid-tempo groove: например, Sade / Rhye / Jungle, а не через резкий танцевальный удар."
        ),
        "dance_block_1": (
            "Первый заход на танцпол: Bruno Mars `Treasure`, Justin Timberlake `CAN'T STOP THE FEELING!`, Dua Lipa `Levitating`, "
            "The Weeknd `Can't Feel My Face`, Earth, Wind & Fire `September`, Jamiroquai `Little L`. "
            "Ритм захода: сначала 2-3 трека, которые дают людям безопасно выйти, не пугать зал EDM-лобовухой."
        ),
        "dance_block_2": (
            "Основной разгон: Purple Disco Machine edits, Dua Lipa, Jamiroquai, Robin S `Show Me Love`, Moloko `Sing It Back`, "
            "ABBA tasteful edits, Calvin Harris без перегруза, Avicii `Levels` только если зал действительно готов. "
            "Логика: поднять sing-along, потом диско-магниты, потом пик."
        ),
        "dance_block_3": (
            "Финальный танцевальный блок собирать по составу гостей. Если зал разновозрастный, держать мост между disco/pop classics и современными edits. "
            "Ориентиры: Whitney Houston `I Wanna Dance with Somebody`, Black Eyed Peas `I Gotta Feeling`, Madonna `Hung Up`, "
            "Ace of Base / Roxette / Руки Вверх только по факту реакции зала и без вкусового провала."
        ),
        "final_block_music": (
            "Под финальные слова: Coldplay `A Sky Full of Stars` (если нужен светлый подъем), "
            "OneRepublic `I Lived`, M83 `Midnight City` soft exit, либо теплый cinematic-pop инструментал. "
            "Финал не должен звучать как разрыв после клуба."
        ),
        "final_music": (
            "Последний трек выбирать так, чтобы люди ушли с ощущением титров хорошего фильма. "
            f"Для {('свадьбы' if is_wedding else 'юбилея')} хорошо работают теплые, объединяющие, узнаваемые композиции без агрессии и без мемности."
        ),
        "stop_list": hard_bans + [
            "джинглы и мемные отбивки",
            "шумовые эффекты поверх речи ведущего",
            "жесткий EDM сразу после трогательного блока",
            "случайная шансонная мешанина без запроса заказчика",
        ],
        "technical_notes": [
            "Перед каждым большим текстом ведущего делать чистый музыкальный коридор минимум 6-8 секунд.",
            "После семейного или эмоционального блока не взрывать зал сразу: сначала warm-up на среднем груве.",
            "Если зал не идет в танец, откатиться на один уровень вниз по энергии и дать людям знакомый хит-магнит.",
            "Не спорить с речью ведущего по громкости. Голос всегда важнее трека.",
            "Переходы между блоками должны звучать как монтаж, а не как случайный autoplay.",
        ],
    }


def extract_music_preferences(questionnaire: dict[str, Any]) -> list[str]:
    return list_from_text(
        "\n".join(
            [
                questionnaire.get("musicLikes", ""),
                questionnaire.get("references", ""),
                questionnaire.get("likedFormats", ""),
                questionnaire.get("jubileeLikedFormats", ""),
            ]
        ),
        [],
    )


def refine_program_payload(program: dict[str, Any], questionnaire: dict[str, Any]) -> dict[str, Any]:
    preferred_tracks = extract_music_preferences(questionnaire)
    timeline = as_list(program.get("scenario_timeline"))
    keep_titles = {
        "Сбор гостей и тонкий welcome",
        "Главное открытие вечера",
        "Первое смысловое ядро",
        "Семейный и эмоциональный блок",
        "Сюрприз или специальный акцент",
        "Танцевальная волна",
        "Финал и закрытие вечера",
    }
    if timeline:
        filtered = [block for block in timeline if as_dict(block).get("block_title") in keep_titles]
        if filtered:
            program["scenario_timeline"] = filtered

    dj = as_dict(program.get("dj_guidance"))
    if preferred_tracks:
        chosen = ", ".join(preferred_tracks[:4])
        dj["overall_music_policy"] = (
            str(dj.get("overall_music_policy", "")).strip()
            + f" Обязательно встроить в сет музыкальные предпочтения клиента: {chosen}."
        ).strip()
        dj["technical_notes"] = as_list(dj.get("technical_notes")) + [
            f"Согласованные любимые треки клиента держать как святые точки вечера: {chosen}.",
            "Если любимый трек медленный, использовать его как эмоциональную точку или красивый переход, а не бросать в случайный блок.",
        ]
        if "special_tracks" not in dj:
            dj["special_tracks"] = preferred_tracks[:6]
        program["dj_guidance"] = dj

    trend_features = [
        "Фишка ведущего: собрать 3 короткие голосовые реплики от гостей заранее и вшить их в эмоциональный блок как живой аудио-монтаж.",
        "Фишка ведущего: использовать micro-story cards — по одной точной истории о паре от близких вместо длинной череды одинаковых поздравлений.",
        "Фишка ведущего: делать мягкий editorial-style guest spotting — коротко замечать стильные, трогательные и характерные детали гостей без кринжа и давления.",
        "Фишка ведущего: собирать танцпол через first circle — сначала вывести 6-8 опорных гостей, а не звать в пустоту весь зал.",
    ]
    program["key_host_commands"] = as_list(program.get("key_host_commands")) + trend_features

    final_print = as_dict(program.get("final_print_version"))
    final_print["timeline_short"] = [
        f"{as_dict(item).get('time_from', '')}-{as_dict(item).get('time_to', '')}: {as_dict(item).get('block_title', '')}"
        for item in as_list(program.get("scenario_timeline"))
    ]
    program["final_print_version"] = final_print
    return program


def build_fallback_program(questionnaire: dict[str, Any]) -> dict[str, Any]:
    event_type = questionnaire["eventType"]
    is_wedding = event_type == "wedding"
    atmosphere = questionnaire.get("atmosphere") or questionnaire.get("anniversaryAtmosphere") or "теплая, живая, современная"
    timing_anchor = questionnaire.get("startTime", "").strip() or "18:00"
    city = questionnaire.get("city", "")
    venue = questionnaire.get("venue", "")
    event_date = questionnaire.get("eventDate", "")
    main_hero = (
        f"{questionnaire.get('groomName', '').strip()} и {questionnaire.get('brideName', '').strip()}".strip()
        if is_wedding
        else questionnaire.get("celebrantName", "").strip()
    )
    key_moments = list_from_text(
        questionnaire.get("keyMoments", ""),
        [
            "сильное открытие без затяжной официальности",
            "эмоциональное ядро вечера",
            "мягкое вовлечение гостей без кринжа",
            "танцевальная волна",
            "чистый, теплый финал",
        ],
    )
    hard_bans = list_from_text(
        questionnaire.get("musicBans", "") + "\n" + questionnaire.get("whatCannotBeDone", ""),
        ["кринжовые конкурсы", "токсичный юмор", "жесткое давление на гостей"],
    )
    sensitive_topics = list_from_text(
        questionnaire.get("conflictTopics", "") + "\n" + questionnaire.get("jubileeConflictTopics", ""),
        ["чувствительные темы без согласования"],
    )

    opening_main = (
        f"Добрый вечер. Сегодня зал собирается не ради формального праздника, а ради живой истории {main_hero}. "
        "Хороший вечер никогда не начинается с шума ради шума. Он начинается с правильного тона: с уважения к людям, "
        "с красивого темпа, с точных слов, которые не давят, а открывают пространство. "
        "Поэтому сегодня мы будем двигаться не по шаблону, а по внутренней логике этого события: там, где нужен свет, дадим свет; "
        "там, где нужна улыбка, не будем заменять ее громкостью; там, где нужна эмоция, не испугаемся тишины."
    )
    family_block_intro = (
        "Есть люди, рядом с которыми любая история становится глубже. Именно они помнят не только яркие кадры, "
        "но и то, что остается за ними: маленькие привычки, тихую поддержку, характер, который не видно на фотографиях. "
        "И сейчас важно дать слово не формальности, а близости."
    )

    return {
        "event_passport": {
            "event_type": event_type,
            "format_name": "Рабочий сценарий для ведущего Event AI",
            "city": city,
            "venue": venue,
            "event_date": event_date,
            "working_timeline_note": "Сценарий собран в fallback-режиме, чтобы система оставалась стабильной и рабочей.",
            "main_goal": "Собрать цельный, современный, персонализированный вечер, пригодный для реальной площадки.",
            "atmosphere": atmosphere,
            "style": "живой, дорогой по ощущению, режиссерски точный, без шаблонной пустоты",
            "mandatory_points": key_moments,
            "hard_bans": hard_bans,
            "timing_anchor": timing_anchor,
        },
        "quality_panel": {
            "scenario_verdict": "Сценарий пригоден как рабочая база для площадки и уже содержит управленческую логику.",
            "director_verdict": "Ритм выстроен волнами: знакомство, раскрытие, эмоциональный пик, энергия, финал.",
            "critic_verdict": "Нужно уточнить несколько деталей по гостям и таймингу кухни, но скелет вечера уже собран профессионально.",
            "final_ready": True,
            "fixed_issues": [
                "Убраны устаревшие конкурсы и токсичные приемы.",
                "Сценарий ориентирован на персонализацию, а не на универсальные штампы.",
                "DJ-блок и риски даны прикладно, а не общими словами.",
            ],
        },
        "concept": {
            "big_idea": f"Вечер, в котором в центре не шаблон праздника, а реальные люди вокруг истории {main_hero}.",
            "main_director_thesis": "Не давить форматом, а вести зал через вкус, внимание к деталям и постепенное раскрытие эмоции.",
            "main_emotional_result": "Гости чувствуют, что были не наблюдателями программы, а частью живого и красивого события.",
            "why_this_event_will_be_remembered": "Потому что здесь работают личные смыслы, точный темп и уважительная, взрослая драматургия вечера.",
        },
        "trend_layer": {
            "trend_summary": "Сценарий строится на современной event-логике: персонализация, режиссура, мягкое вовлечение и музыкальная драматургия.",
            "applied_trends": [
                "отказ от кринжовых активностей в пользу точных персональных заходов",
                "сильные, но не крикливые речевые блоки ведущего",
                "гибкий темп с музыкальными волнами вместо монотонного застолья",
                "бережная работа с чувствительными темами и уместностью гостей",
            ],
            "rejected_outdated_patterns": [
                "навязчивые конкурсы ради заполнения пауз",
                "шаблонные тосты без привязки к людям",
                "свадебные и юбилейные штампы, звучащие как чужой текст",
            ],
        },
        "key_host_commands": [
            f"Не трогать без необходимости темы: {', '.join(sensitive_topics)}.",
            "Первыми поднимать самых открытых и поддерживающих атмосферу гостей.",
            "Не жертвовать эмоциональным ядром ради случайного шумного интерактива.",
            "Святые точки вечера заранее проговорить с заказчиком и держать их в тайминге до последнего.",
            "Если зал тяжелый на старт, сначала собирать доверие, а не форсировать веселье.",
        ],
        "questions_to_clarify_before_event": [
            "Есть ли точный тайминг горячего, торта, сюрпризов и технических включений?",
            "Кого точно нельзя вовлекать в активные блоки?",
            "Есть ли обязательные поздравления по списку и каков приоритет этих людей?",
            "Есть ли сюрпризы от гостей или семьи, которые нужно встроить заранее?",
        ],
        "director_logic": {
            "opening_logic": "Открытие строится через доверие и красивое включение зала, без шума ради шума.",
            "development_logic": "Вечер развивается через чередование смысловых точек, гостевого участия и музыкальных волн.",
            "family_or_core_emotional_logic": "Эмоциональное ядро ставится тогда, когда зал уже слушает и готов переживать, а не только наблюдать.",
            "final_logic": "Финал не растягивается, а собирает вечер в ясную, теплую и запоминающуюся точку.",
        },
        "scenario_timeline": [
            {
                "time_from": timing_anchor,
                "time_to": "18:20",
                "block_title": "Сбор гостей и мягкий welcome",
                "block_purpose": "Снять стартовое напряжение и собрать зал в единое поле внимания.",
                "what_happens": "Гости заходят, рассаживаются, атмосфера собирается через музыку и спокойное присутствие ведущего.",
                "host_action": "Ведущий приветствует, ориентирует и аккуратно задает тон вечера.",
                "host_text": "Первое впечатление о вечере складывается раньше, чем прозвучат главные слова. Поэтому старт этого события должен быть не суетливым, а красивым: с воздухом, с ощущением заботы, с внутренним спокойствием. Зал еще только собирается, а мы уже закладываем главный принцип вечера: здесь никто не будет гнаться за шумом, здесь будут создавать настроение тоньше и точнее.",
                "dj_task": "Держать стильный welcome без резких переходов и громких вокальных пиков.",
                "director_move": "Сначала атмосфера, потом акценты.",
                "risk_control": "Не перегружать зал объявлениями, пока гости физически не собраны.",
                "transition": "Когда в зале большинство гостей и внимание стабилизируется, перевести пространство к основному открытию.",
            },
            {
                "time_from": "18:20",
                "time_to": "18:40",
                "block_title": "Главное открытие",
                "block_purpose": "Задать характер вечера, стиль общения и эмоциональный вектор.",
                "what_happens": "Ведущий открывает событие, обозначает интонацию и внутреннюю логику вечера.",
                "host_action": "Говорить образно, но точно, без банальностей и лишнего пафоса.",
                "host_text": opening_main,
                "dj_task": "Короткий кинематографичный вход, затем чистый музыкальный уход под речь.",
                "director_move": "Сформировать доверие к ведущему и желание слушать дальше.",
                "risk_control": "Если зал шумный, использовать короткую версию открытия и вернуться к большой речи позже.",
                "transition": "После открытия перевести внимание к первой содержательной точке вечера.",
            },
            {
                "time_from": "18:40",
                "time_to": "19:20",
                "block_title": "Первое смысловое ядро",
                "block_purpose": "Дать залу не только информацию, но и эмоциональную глубину.",
                "what_happens": "Через близких, историю или важные факты вечер получает человеческий объем и персональную основу.",
                "host_action": "Управлять темпом речей и не превращать теплый блок в затянутую официальность.",
                "host_text": family_block_intro,
                "dj_task": "Мягкие инструментальные подложки, поддерживающие речь, а не спорящие с ней.",
                "director_move": "Углубить зал, прежде чем поднимать энергию.",
                "risk_control": "Если блок затягивается, аккуратно резать количество выходов, а не растягивать эмоцию до усталости.",
                "transition": "После теплой глубины вывести зал в более живое движение через легкий интерактив или музыкальный переход.",
            },
            {
                "time_from": "19:20",
                "time_to": "20:00",
                "block_title": "Мягкое вовлечение гостей",
                "block_purpose": "Поднять энергию, не ломая достоинство вечера.",
                "what_happens": "Работа с гостями строится через ассоциации, личные истории и уместные точечные включения.",
                "host_action": "Поднимать только тех, кто подходит по энергии и уместности; не вытаскивать людей насильно.",
                "host_text": "Самый хороший интерактив это тот, после которого гости чувствуют не неловкость, а удовольствие от того, что их увидели и услышали. Поэтому сейчас важна не форма ради формы, а точность: кого вовлечь, в каком объеме и в какой интонации.",
                "dj_task": "Дать короткие отбивки и музыкальные мосты, не скатываясь в мемность и клоунаду.",
                "director_move": "Переключить зал из слушания в живое присутствие.",
                "risk_control": "Стоп-лист гостей и тем держать в голове постоянно.",
                "transition": "На хорошем уровне энергии перевести вечер к танцу, гастрономической паузе или следующему смысловому блоку.",
            },
            {
                "time_from": "20:00",
                "time_to": "21:00",
                "block_title": "Танцевальная волна и энергетический пик",
                "block_purpose": "Дать залу телесную энергию и ощущение праздника без потери вкуса.",
                "what_happens": "Через выстроенную музыкальную волну гости включаются в движение, а вечер выходит на пик живости.",
                "host_action": "Не мешать музыке лишним микрофоном, а запускать танцпол короткими точными командами.",
                "host_text": "Есть моменты, которые не нужно подробно объяснять. Их нужно просто вовремя отпустить в музыку. И если до этого вечер раскрывался словами и смыслами, то сейчас он должен задышать шире, свободнее, телеснее. Именно так праздник становится не просто красивым, а по-настоящему прожитым.",
                "dj_task": "Начать с комфортного узнавания, затем поднять зал на более яркую, но не безвкусную энергию.",
                "director_move": "Пик должен ощущаться как естественный рост, а не как внезапная попытка раскачать зал.",
                "risk_control": "Если танцпол не идет, сначала дать более знакомый материал средней энергии, а не сразу бить в максимальную громкость.",
                "transition": "После активной волны вернуть внимание к финальному смысловому завершению.",
            },
            {
                "time_from": "21:00",
                "time_to": "21:30",
                "block_title": "Финальное эмоциональное собирание",
                "block_purpose": "Собрать вечер в чистую и запоминающуюся финальную точку.",
                "what_happens": "Финальные слова, важный общий момент, при необходимости общий жест, торт или завершающая музыкальная композиция.",
                "host_action": "Говорить коротко, сильно и тепло, не растягивая финал.",
                "host_text": "Хороший вечер заканчивается не в тот момент, когда выключают музыку, а в тот, когда у людей внутри остается ясное чувство: это было не случайное собрание, а настоящее событие. И если сегодня каждый унесет с собой это ощущение, значит все было не зря.",
                "dj_task": "Подложить объединяющую композицию и аккуратно собрать зал в финальный эмоциональный кадр.",
                "director_move": "Завершить мягко, но собранно, без провисания.",
                "risk_control": "Если гости устали, убрать лишние слова и не превращать финал в второй финал.",
                "transition": "После официального завершения оставить залу свободный мягкий хвост вечера.",
            },
        ],
        "host_script": {
            "opening_main": opening_main,
            "opening_short": "Добрый вечер. Пусть это событие с первых минут будет красивым, точным и по-настоящему вашим.",
            "welcome_line": "Пусть сегодня все начнется не с суеты, а с ощущения, что вы попали именно туда, где должны быть.",
            "first_core_intro": "Прежде чем идти дальше по ритму вечера, важно открыть ту часть истории, без которой все остальное было бы просто внешней формой.",
            "family_block_intro": family_block_intro,
            "surprise_intro": "Следующий момент важен не как эффект ради эффекта, а как живая эмоция, которая останется после праздника дольше любой декорации.",
            "dance_block_intro": "А теперь вечеру пора сменить походку: от красивого разговора перейти в музыку, движение и свободную энергию зала.",
            "final_block_intro": "Перед тем как отпустить этот вечер дальше, хочется собрать его в одну ясную и честную мысль.",
            "closing_words": "Спасибо за атмосферу, за внимание друг к другу и за то, что этот вечер получился не громким, а настоящим. Именно такие события люди вспоминают дольше всего.",
        },
        "dj_guidance": {
            "overall_music_policy": "Музыкальная линия должна расти волнами: стильный welcome, чистое открытие, аккуратная поддержка смысловых сцен, затем уверенный танцевальный подъем и теплый финал.",
            "welcome_music": "Nu-disco, soft funk, elegant pop, lounge-pop, современный легкий groove без навязчивого давления вокалом.",
            "opening_music": "Короткий кинематографичный старт с мягким спадом под голос ведущего.",
            "table_background": "Soul-pop, soft disco, легкий funk, негромкий ритм, чтобы не спорить с речами гостей.",
            "emotional_blocks_music": "Piano, strings, atmospheric pads, инструментальные версии без чрезмерной сентиментальности.",
            "dance_block_1": "Поднять через знакомые треки средней энергии, чтобы зал не испугался входа в танец.",
            "dance_block_2": "Основной пик: уверенные sing-along треки, стильные disco edits, узнаваемые хиты без безвкусной мешанины.",
            "dance_block_3": "Финальный блок подбирать по возрасту и составу гостей, сохраняя вкус и энергию до конца.",
            "final_block_music": "Теплая объединяющая композиция без агрессивного бита, чтобы финал был красивым, а не шумным.",
            "final_music": "Финальный трек должен оставлять ощущение завершенности, а не провоцировать хаотичное добивание танцпола.",
            "stop_list": hard_bans,
            "technical_notes": [
                "Не перебивать ведущего джинглами и мемными звуками.",
                "На эмоциональных блоках держать запас по громкости.",
                "Перед выходами гостей давать чистый музыкальный коридор.",
                "После трогательных блоков поднимать энергию постепенно, а не резким скачком.",
            ],
        },
        "guest_management": {
            "active_people": ["Первыми работать с самыми открытыми и поддерживающими атмосферу гостями."],
            "shy_people": ["Не выводить в центр без предварительного контакта и явного согласия."],
            "important_people": ["Родители, близкие родственники, люди с реальной эмоциональной ценностью для вечера."],
            "do_not_involve": sensitive_topics,
            "sensitive_people_or_topics": sensitive_topics,
            "management_notes": [
                "Не превращать персональную историю в публичный допрос.",
                "Балансировать внимание между героями вечера и залом.",
                "Каждый гость должен чувствовать уважение, даже если он не включен активно.",
            ],
        },
        "risk_map": [
            {
                "risk": "Провал по динамике в первой трети вечера",
                "why_it_matters": "Если старт вялый, зал позже сложнее собрать в единый ритм.",
                "how_to_prevent": "Короткое сильное открытие, быстрая сборка внимания и ранняя персональная точка.",
                "what_to_do_if_triggered": "Сократить длинные речи и раньше перевести вечер в живой формат включения гостей.",
            },
            {
                "risk": "Неудачное касание чувствительной темы",
                "why_it_matters": "Это способно резко разрушить доверие к ведущему и атмосферу события.",
                "how_to_prevent": "Заранее собрать стоп-лист и не импровизировать на спорных темах.",
                "what_to_do_if_triggered": "Быстро увести фокус на нейтральный теплый блок без оправданий и долгих комментариев.",
            },
            {
                "risk": "Танцпол не включается с первого захода",
                "why_it_matters": "Слишком жесткий вход в танец может посадить энергию зала еще сильнее.",
                "how_to_prevent": "Начинать с узнаваемого материала средней энергии и правильно подводить зал речью.",
                "what_to_do_if_triggered": "Сделать короткий откат, дать еще одну смысловую точку и вернуться к танцу через 7-10 минут.",
            },
        ],
        "plan_b": [
            {
                "situation": "Гости долго рассаживаются или задерживается старт",
                "solution": "Продлить welcome, укоротить первое официальное открытие и сместить основной смысловой блок, не ломая общую драматургию.",
            },
            {
                "situation": "Поздравительные речи затягиваются",
                "solution": "Ограничить количество живых выходов, часть сообщений перевести в более компактный формат или перенести после следующей волны энергии.",
            },
            {
                "situation": "Техника или музыка дают сбой",
                "solution": "Ведущий уходит в короткий живой текст и ручное удержание внимания, пока команда тихо чинит техническую часть.",
            },
        ],
        "final_print_version": {
            "title": f"Краткая версия: {main_hero or 'событие'}",
            "summary": "Современный, персонализированный, режиссерски собранный вечер с фокусом на уместность, вкус и живую атмосферу.",
            "timeline_short": [
                f"{timing_anchor} — сбор гостей и welcome",
                "главное открытие",
                "первое смысловое ядро",
                "мягкое вовлечение гостей",
                "танцевальная волна",
                "финальное собирание",
            ],
            "must_do": [
                "держать ритм волнами, а не сплошным шумом",
                "работать через личные смыслы и точные формулировки",
                "согласовать критические точки вечера заранее",
            ],
            "must_not_do": hard_bans,
            "host_focus": [
                "доверие зала",
                "бережность к людям",
                "точность тона",
                "контроль тайминга ключевых блоков",
            ],
            "dj_focus": [
                "чистые музыкальные коридоры",
                "поддержка эмоциональных сцен",
                "безопасный и вкусный разгон танцпола",
            ],
        },
    }


def build_target_program(questionnaire: dict[str, Any]) -> dict[str, Any]:
    program = build_fallback_program(questionnaire)
    event_type = questionnaire["eventType"]
    is_wedding = event_type == "wedding"
    atmosphere = questionnaire.get("atmosphere") or questionnaire.get("anniversaryAtmosphere") or "теплая, живая, современная"
    trend_bank = build_trend_bank(questionnaire)
    style_bank = build_style_bank(questionnaire)
    heroes = (
        f"{questionnaire.get('groomName', '').strip()} и {questionnaire.get('brideName', '').strip()}".strip()
        if is_wedding
        else questionnaire.get("celebrantName", "").strip()
    ) or questionnaire.get("clientName", "").strip() or "герои вечера"
    direct_host_texts = build_direct_host_texts(questionnaire, heroes, atmosphere)
    hard_bans = as_list(as_dict(program.get("event_passport")).get("hard_bans")) or [
        "кринжовые конкурсы",
        "давление на гостей",
        "токсичный юмор",
    ]

    program["event_passport"]["format_name"] = "Подробный рабочий сценарий для ведущего"
    program["event_passport"]["working_timeline_note"] = "Документ собран как прикладной сценарий для реальной площадки: с логикой, рисками, таймингом, текстами ведущего и DJ-блоком."
    program["event_passport"]["main_goal"] = "Выдать ведущему не обзор анкеты, а подробный, современный, персонализированный сценарий, который можно брать в работу без шаблонной пустоты."
    program["event_passport"]["atmosphere"] = atmosphere
    program["event_passport"]["style"] = "живой, дорогой по ощущению, сценичный, современный, без кринжа и без банкетных штампов"
    program["quality_panel"]["scenario_verdict"] = "Сценарий обязан работать как постановочный документ: ведущий, DJ и команда понимают, что делать, зачем и в какой последовательности."
    program["quality_panel"]["director_verdict"] = "Внутри сценария есть режиссерская арка: открытие, развитие, эмоциональное ядро, управляемый подъем энергии и финальное закрытие."
    program["quality_panel"]["critic_verdict"] = "Если в анкете не хватает деталей, они вынесены в уточнения, но сценарий все равно остается рабочим и не разваливается."
    program["quality_panel"]["fixed_issues"] = [
        "Убран обзорный стиль и добавлен прикладной тайминг по блокам.",
        "Тексты ведущего увеличены и переведены в сценичную, метафоричную манеру.",
        "DJ-блок стал инструментом для работы, а не общим пожеланием.",
        "Финальная печатная версия заточена под площадку, а не под презентацию.",
        f"Trend-bank обновлен под {trend_bank['event_year']} год и сезон {trend_bank['season']}.",
    ]
    program["concept"]["big_idea"] = f"{heroes} в центре не как формальный повод, а как живая история, вокруг которой выстраивается красивый, взрослый, современный {('свадебный' if is_wedding else 'юбилейный')} вечер."
    program["concept"]["main_director_thesis"] = "Вести зал не шумом и не шаблоном, а точной драматургией: через личный смысл, музыкальные волны, уважение к гостям и ощутимый режиссерский вкус."
    program["concept"]["main_emotional_result"] = "Гости должны не просто вспомнить программу, а почувствовать, что прожили цельный, красивый и личный вечер."
    program["concept"]["why_this_event_will_be_remembered"] = "Потому что сценарий опирается на конкретных людей, аккуратную режиссуру, сильные тексты ведущего и живую музыкальную логику, а не на устаревшие шаблоны."
    program["trend_layer"]["trend_summary"] = (
        f"Premium-режим: сценарий собирается с учетом trend-bank для {trend_bank['event_year']} года и сезона {trend_bank['season']}. "
        f"Рыночная логика: {trend_bank['seasonal_logic']}."
    )
    program["trend_layer"]["applied_trends"] = as_list(trend_bank["applied_market_direction"]) + [
        "персонализация через реальные детали героев вечера, а не универсальные формулы",
        "режиссерская сборка вечера волнами вместо однотонного застолья",
        "мягкое вовлечение гостей без давления и без кринжовых конкурсов",
        "музыкальная драматургия как часть сценария, а не отдельная функция DJ",
        f"style-bank: {style_bank['voice_direction']}",
    ]
    program["trend_layer"]["rejected_outdated_patterns"] = [
        "конкурсы ради шума",
        "банкетные клише и сладкая пустая романтика",
        "токсичный юмор и давление на гостей",
        "случайный разгон танцпола без режиссерской логики",
    ]
    program["key_host_commands"] = [
        "Не превращать личные истории в пересказ анкеты; каждая деталь должна работать на эмоцию и драматургию.",
        "Не поднимать в актив без согласия скромных, возрастных или чувствительных гостей.",
        "Первыми на вовлечение выводить самых теплых и поддерживающих людей, которые задают безопасную модель поведения залу.",
        "Святые точки вечера держать в тайминге до последнего: открытие, эмоциональное ядро, специальный акцент, финал.",
        "Не жертвовать семейным или смысловым блоком ради случайной суеты на танцполе.",
        "Если зал тяжелый, сначала строить доверие, потом энергию; не наоборот.",
    ]
    program["questions_to_clarify_before_event"] = [
        "Кто из гостей точно должен получить слово, а кого лучше не выводить в микрофонный фокус?",
        "Есть ли сюрпризы, видео, номера или специальные выходы, которые пока не попали в анкету?",
        "Какой реальный тайминг кухни, торта, первого танца, артистов и технических включений?",
        "Какие темы, имена, семейные детали и шутки допустимы только в узком круге, но не на весь зал?",
        "Кто из гостей может стать опорой ведущего в момент вовлечения зала?",
    ]
    program["director_logic"] = {
        "opening_logic": "Открытие должно быстро создать уважительную тишину, ощущение стиля и чувство, что вечер ведет профессиональная рука, а не поток случайных реплик.",
        "development_logic": "Дальше вечер идет волнами: личный смысл, эмоция, облегчение, вовлечение, акцент, энергия, финал. Каждая следующая точка вытекает из предыдущей.",
        "family_or_core_emotional_logic": "Эмоциональное ядро ставится тогда, когда зал уже доверяет ведущему и готов не просто смотреть, а переживать вместе с героями вечера.",
        "final_logic": "Финал не должен быть долгим. Он должен быть точным, собранным и теплым, чтобы гости унесли с собой не шум, а послевкусие.",
    }
    program["scenario_timeline"] = build_detailed_timeline(questionnaire, heroes, atmosphere, hard_bans)
    program["host_script"] = build_detailed_host_script(questionnaire, heroes, atmosphere)
    program["dj_guidance"] = {
        "overall_music_policy": "Музыка работает как драматургия, а не как фон из случайных треков. Сначала собираем доверие и стиль, потом двигаем энергию, а эмоциональные блоки защищаем от грубых музыкальных вторжений.",
        "welcome_music": "Nu-disco, soul-pop, elegant house, lounge edits с мягким грувом. Нужен звук дорогого старта, а не банкетной суеты.",
        "opening_music": "Кинематографичный инструментал или атмосферный electronic-pop intro с мягким спадом под голос ведущего.",
        "table_background": "Легкий groove, soul, tasteful pop, clean funk. Музыка должна поддерживать разговоры, а не спорить с ними.",
        "emotional_blocks_music": "Piano, strings, ambient textures, не слишком сладкие инструменталы. После эмоционального пика выходить вверх плавно, без резкого удара барабана.",
        "dance_block_1": "Старт танцпола через узнаваемые треки средней энергии, которые дают людям чувство безопасности и желания присоединиться.",
        "dance_block_2": "Главный разгон через sing-along, disco edits, pop-dance и треки-магниты без дешевой мешанины и без ломаного жанрового хаоса.",
        "dance_block_3": "Финальная энергия строится по составу гостей: держать вкус, избегать крика, не уходить в случайный набор ради количества.",
        "final_block_music": "Объединяющий красивый трек с теплым эмоциональным кодом, под который можно дать финальные слова и общий кадр.",
        "final_music": "Последний трек должен закрывать вечер как титры хорошего фильма: с завершенностью, а не с ощущением недожатого хаоса.",
        "stop_list": hard_bans + ["мемные джинглы", "резкие обрубания подводок", "треки, ломающие трогательные сцены"],
        "technical_notes": [
            "Не перебивать ведущего джинглами и звуковыми шутками.",
            "Перед речью всегда оставлять чистый музыкальный коридор.",
            "На эмоциональных сценах держать запас по громкости и не спорить с интонацией голоса.",
            "После трогательных моментов поднимать энергию ступенчато, а не прыжком.",
            "Если танцпол не идет, откатиться на один уровень по энергии и снова собрать зал знакомым материалом.",
        ],
    }
    for block in program["scenario_timeline"]:
        title = block.get("block_title", "")
        if title in direct_host_texts:
            block["host_text"] = direct_host_texts[title]
    program["host_script"] = {
        "opening_main": direct_host_texts["Главное открытие вечера"],
        "opening_short": direct_host_texts["Главное открытие вечера"][:900],
        "welcome_line": direct_host_texts["Сбор гостей и тонкий welcome"],
        "first_core_intro": direct_host_texts["Первое смысловое ядро"],
        "family_block_intro": direct_host_texts["Семейный и эмоциональный блок"],
        "surprise_intro": direct_host_texts["Сюрприз или специальный акцент"],
        "dance_block_intro": direct_host_texts["Танцевальная волна"],
        "final_block_intro": direct_host_texts["Финал и закрытие вечера"],
        "closing_words": direct_host_texts["Финал и закрытие вечера"],
    }
    program["dj_guidance"] = build_precise_dj_guidance(questionnaire, hard_bans)
    program["guest_management"] = {
        "active_people": ["Вовлекать первыми тех, кто улыбается, поддерживает атмосферу и не боится публичного внимания."],
        "shy_people": ["Не вытягивать в центр молчаливых и скромных гостей без предварительного контакта и явного согласия."],
        "important_people": ["Родители, близкие, люди с реальным эмоциональным весом для героев вечера должны быть встроены в сценарий заранее, а не по импровизации."],
        "do_not_involve": ["Закрытых, уставших, конфликтных и явно неготовых к публичности гостей не вовлекать в активные блоки."],
        "sensitive_people_or_topics": ["Разводы, конфликты, потери, финансовые темы, неловкие шутки о возрасте, внешности, детях или прошлом без прямого согласования недопустимы."],
        "management_notes": [
            "Гостями нужно управлять через уважение и правильную последовательность, а не через давление.",
            "Если зал устал, сначала вернуть внимание смыслом и теплом, а не форсировать громкость.",
            "Каждый гость должен чувствовать, что его не используют как реквизит ради активности.",
        ],
    }
    program["risk_map"] = [
        {
            "risk": "Бедное или шаблонное открытие",
            "why_it_matters": "Если вечер стартует без статуса и образа, дальше любое вовлечение будет ощущаться дешевле, чем должно.",
            "how_to_prevent": "Открывать вечер сильной метафоричной речью, коротко, точно и с ясной режиссерской рамкой.",
            "what_to_do_if_triggered": "Остановить суету, вернуть тишину, дать короткий сильный текст и заново собрать внимание зала.",
        },
        {
            "risk": "Провисание середины вечера",
            "why_it_matters": "Средняя часть определяет, запомнится ли вечер как цельная история или как набор отдельных действий.",
            "how_to_prevent": "Чередовать личные смыслы, гостевое участие, музыкальные волны и специальные акценты.",
            "what_to_do_if_triggered": "Укоротить проходные реплики, быстрее выйти к следующему смысловому или энергетическому блоку.",
        },
        {
            "risk": "Кринжовое вовлечение гостей",
            "why_it_matters": "Оно моментально сбивает статус ведущего и разрушает доверие к сценарию.",
            "how_to_prevent": "Использовать мягкие адресные включения, наблюдательность и уважительную коммуникацию вместо конкурсов ради конкурсов.",
            "what_to_do_if_triggered": "Сразу снять давление шуткой на себя, вернуть зал к безопасной теме и переформатировать активность.",
        },
        {
            "risk": "Неровный музыкальный разгон",
            "why_it_matters": "Слишком резкий вход в танец или плохой выход из трогательного блока ломает драматургию вечера.",
            "how_to_prevent": "Собирать музыкальные волны ступенчато и держать DJ в тесной связке с ведущим.",
            "what_to_do_if_triggered": "Сделать короткий откат по энергии, заново собрать внимание и только потом снова поднимать зал.",
        },
    ]
    program["plan_b"] = [
        {"situation": "Задерживается сбор гостей или кухня", "solution": "Продлить welcome и укоротить первое официальное включение, сохранив основную драматургию."},
        {"situation": "Слишком затянулись поздравления", "solution": "Часть речей перевести в более короткий формат и перенести акцент на следующую смысловую точку."},
        {"situation": "Сюрприз срывается или технически не готов", "solution": "Ведущий закрывает паузу живым текстом, а команда переставляет специальный акцент на позже."},
        {"situation": "Танцпол не включился с первого захода", "solution": "Вернуться к знакомому треку средней энергии, дать короткую подводку и собрать опорных гостей в первый круг."},
    ]
    program["final_print_version"] = {
        "title": f"Краткая печатная версия: {heroes}",
        "summary": "Подробный рабочий сценарий с режиссерской логикой, длинными текстами ведущего, прикладным DJ-блоком и понятным планом действий на площадке.",
        "timeline_short": [f"{item['time_from']}-{item['time_to']}: {item['block_title']}" for item in program["scenario_timeline"]],
        "must_do": [
            "Держать вечер волнами, а не сплошным шумом.",
            "Опирайтесь на личные смыслы героев, а не на универсальные банкетные штампы.",
            "Защищайте эмоциональное ядро и финал от случайной суеты.",
            "Работайте с DJ как с соавтором драматургии.",
        ],
        "must_not_do": hard_bans + ["не скатываться в банальности", "не жертвовать смыслом ради активности"],
        "host_focus": ["точность тона", "уважение к людям", "контроль тайминга", "мягкое вовлечение", "сильные образные тексты"],
        "dj_focus": ["музыкальные коридоры", "бережный выход из эмоции", "пошаговый разгон", "чистый финальный акцент"],
    }
    return refine_program_payload(program, questionnaire)


def is_program_detailed(program: dict[str, Any]) -> bool:
    timeline = as_list(program.get("scenario_timeline"))
    if len(timeline) < 7:
        return False
    for block in timeline:
        block_data = as_dict(block)
        if len(str(block_data.get("host_text", "")).strip()) < 900:
            return False
    script = as_dict(program.get("host_script"))
    for key in [
        "opening_main",
        "first_core_intro",
        "family_block_intro",
        "dance_block_intro",
        "closing_words",
    ]:
        if len(str(script.get(key, "")).strip()) < 900:
            return False
    return True


def build_event_analyst_brief(questionnaire: dict[str, Any]) -> dict[str, Any]:
    event_type = questionnaire.get("eventType", "")
    lead_heroes = (
        f"{questionnaire.get('groomName', '').strip()} и {questionnaire.get('brideName', '').strip()}".strip()
        if event_type == "wedding"
        else questionnaire.get("celebrantName", "").strip()
    )
    atmosphere = questionnaire.get("atmosphere") or questionnaire.get("anniversaryAtmosphere") or "современная, живая, собранная"
    sensitive = list_from_text(
        questionnaire.get("conflictTopics", "") + "\n" + questionnaire.get("jubileeConflictTopics", ""),
        ["избегать чувствительных тем без согласования"],
    )
    personalization = build_personalization_brief(questionnaire)
    return {
        "event_type": event_type,
        "lead_heroes": lead_heroes or questionnaire.get("clientName", "").strip() or "герои вечера",
        "city": questionnaire.get("city", ""),
        "venue": questionnaire.get("venue", ""),
        "date": questionnaire.get("eventDate", ""),
        "start_time": questionnaire.get("startTime", ""),
        "atmosphere": atmosphere,
        "goal": "выдать пригодный для площадки рабочий сценарий, а не обзор анкеты",
        "sensitive_topics": sensitive,
        "must_keep": list_from_text(questionnaire.get("keyMoments", ""), ["сильное открытие", "эмоциональное ядро", "финал"]),
        "music_constraints": list_from_text(questionnaire.get("musicBans", ""), ["не ломать эмоциональные сцены"]),
        "personalization_mandates": {
            "love_story": personalization["love_story"],
            "proposal_story": personalization["proposal_story"],
            "values": personalization["values"],
            "nicknames": personalization["nicknames"] or "использовать личный язык пары, если он уместен",
            "inside_jokes": personalization["inside_jokes"] or "не придумывать шутки с нуля, а опираться на реальные внутренние коды пары",
            "important_dates": personalization["important_dates"] or ["вшить реальные важные точки истории пары"],
            "host_wishes": personalization["host_wishes"] or "соблюсти пожелания к тону ведущего",
        },
    }


def build_trend_analyst_brief(questionnaire: dict[str, Any]) -> dict[str, Any]:
    if questionnaire.get("eventType") == "wedding":
        return {
            "tone": "editorial romance without sugar",
            "hosting_style": [
                "образная сценичная речь без свадебных штампов",
                "мягкое вовлечение вместо конкурсов",
                "ощущение дорогой режиссуры и вкуса",
            ],
            "music_logic": [
                "welcome через nu-disco, soul-pop, elegant house",
                "эмоциональные сцены через piano, strings, ambient textures",
                "танцпол через узнаваемые pop-dance и disco edits без вкусового провала",
            ],
        }
    return {
        "tone": "warm prestige with wit and biography-driven dramaturgy",
        "hosting_style": [
            "сценичная живая речь без канцелярита",
            "уважение к возрасту и статусу гостей без официоза",
            "ирония и свет вместо пафоса",
        ],
        "music_logic": [
            "welcome через soul, lounge, light funk",
            "эмоциональные блоки без сладкой ретро-банальности",
            "танцпол через возрастной микс с вкусом и понятным разгоном",
        ],
    }


def build_system_prompt() -> str:
    return """
Ты — сильный event-режиссер, сценарист ведущих и редактор финального рабочего документа.

Работаешь только с двумя типами мероприятий:
- wedding
- jubilee

Твоя задача: на основе анкеты собрать ГОТОВЫЙ рабочий сценарий для ведущего.
Это не обзор анкеты и не список советов.
Это должен быть прикладной документ, пригодный для настоящей работы на площадке.

Критические требования:
- один основной сильный вызов
- никакой воды
- никакого канцелярита
- никакой шаблонной романтической или юбилейной банальности
- не предлагай кринжовые конкурсы, токсичный юмор и устаревшие форматы
- стиль речи ведущего должен быть образным, метафоричным, сценичным, но пригодным к чтению вслух
- ключевые тексты ведущего должны быть длинными и содержательными
- DJ guidance должен быть прикладным, с музыкальной логикой, а не абстракцией
- если данных не хватает, не ломайся: строй лучший рабочий вариант и отдельно выноси уточнения

Верни строго JSON со следующими ключами:
event_passport
quality_panel
concept
trend_layer
key_host_commands
questions_to_clarify_before_event
director_logic
scenario_timeline
host_script
dj_guidance
guest_management
risk_map
plan_b
final_print_version
"""


def build_stage_system_prompt() -> str:
    return """
Ты работаешь как одна быстрая, но сильная связка ролей:
- сценарист-драматург
- trend analyst
- style editor
- креативный копирайтер
- DJ editor
- режиссер
- критик
- финальная сборка

Работаешь только с двумя типами мероприятий:
- wedding
- jubilee

Твоя задача: на основе анкеты собрать ГОТОВЫЙ рабочий сценарий для ведущего.
Это не обзор анкеты и не список советов.
Это прикладной документ для настоящей работы на площадке.
Тебе передаются trend-bank, style-bank и dramaturgy-bank. Они обязательны к использованию.
Если дата события указывает на будущий сезон или год, ты ориентируешься на них, а не на устаревшую общую логику.

Критические требования:
- один основной сильный AI-вызов
- никакой воды
- никакого канцелярита
- никакой шаблонной романтической или юбилейной банальности
- никаких кринжовых конкурсов, токсичного юмора и устаревших форматов
- стиль речи ведущего должен быть образным, метафоричным, небанальным, остроумным, сценичным, живым и эмоционально гибким
- метафоры должны быть свежими, а не заезженными
- текст должен звучать как авторская речь для реальной сцены, а не как генератор открыток
- нельзя использовать повторяющиеся заготовки, взаимозаменяемые абзацы и универсальные формулы, которые можно переставить в другую анкету без потери смысла
- каждая программа должна писаться заново под конкретную анкету, а не быть вариацией одного и того же шаблона
- запрещено механически пересказывать слова анкеты; сначала сделай выводы, потом пиши
- сначала анализируй людей, динамику пары, риск-карту и энергетику вечера, и только потом собирай тексты
- воспринимай анкету не как техническое задание, а как источник вдохновения и драматургической фактуры
- используй технику метафорического переноса: найди центральную метафору пары и построй вокруг нее образную логику вечера
- работай по формуле detail -> meaning -> stage delivery
- соблюдай negative prompt lexicon из style-bank и не используй стоп-слова
- ключевые тексты ведущего должны быть длинными и пригодными к чтению 3-5 минут
- host_text в каждом блоке тайминга должен содержать не менее 10 интересных метафорических образов
- DJ guidance должен быть прикладным: трековые ориентиры, ритм захода, связки, что не включать, как поднимать зал, как спасать просадку и как не ломать эмоциональные сцены
- если данных не хватает, не ломайся: строй лучший рабочий вариант и отдельно выноси уточнения

Как мыслить внутри одного ответа:
1. Dramaturg: найди центральную метафору пары, уникальную зацепку и 2-3 вспомогательных образа.
2. Trend analyst: используй trend-bank с учетом года и даты события; отфильтруй устаревшие приемы и определи актуальную event-логику.
3. Style editor: используй style-bank, чтобы отрезать банальность, пересказ анкеты, стоп-слова и взаимозаменяемые абзацы.
4. Copywriter: преврати детали анкеты в смысл, образы и готовые сценические реплики.
5. Сценарист: построй подробный тайминг и длинные тексты ведущего.
6. DJ editor: собери реальный DJ sheet по музыкальным точкам вечера и любимым трекам клиента.
7. Режиссер: проверь ритм, переходы, музыкальную логику и управление залом.
8. Критик: найди слабые места и поправь их до финальной версии.
9. Финальная сборка: перед возвратом молча проверь каждый ключевой блок и перепиши любой фрагмент, который звучит как универсальная заготовка, как механический пересказ анкеты или как канцелярит.
10. Верни только сильный итоговый JSON.

Верни строго JSON с ключами:
event_passport
quality_panel
concept
trend_layer
key_host_commands
questions_to_clarify_before_event
director_logic
scenario_timeline
host_script
dj_guidance
guest_management
risk_map
plan_b
final_print_version
"""


def build_generation_user_prompt(questionnaire: dict[str, Any]) -> str:
    analyst_brief = json.dumps(build_event_analyst_brief(questionnaire), ensure_ascii=False, indent=2)
    trend_brief = json.dumps(build_trend_analyst_brief(questionnaire), ensure_ascii=False, indent=2)
    personalization_brief = json.dumps(build_personalization_brief(questionnaire), ensure_ascii=False, indent=2)
    trend_bank = json.dumps(build_trend_bank(questionnaire), ensure_ascii=False, indent=2)
    style_bank = json.dumps(build_style_bank(questionnaire), ensure_ascii=False, indent=2)
    dramaturgy_bank = json.dumps(build_dramaturgy_bank(questionnaire), ensure_ascii=False, indent=2)
    return (
        f"Тип мероприятия: {questionnaire['eventType']}\n\n"
        "Анкета:\n"
        f"{build_questionnaire_context(questionnaire)}\n\n"
        "Бриф аналитика:\n"
        f"{analyst_brief}\n\n"
        "Бриф trend analyst:\n"
        f"{trend_brief}\n\n"
        "Персональный профиль героев:\n"
        f"{personalization_brief}\n\n"
        "Trend-bank:\n"
        f"{trend_bank}\n\n"
        "Style-bank:\n"
        f"{style_bank}\n\n"
        "Dramaturgy-bank:\n"
        f"{dramaturgy_bank}\n\n"
        "Верни только валидный JSON. Нельзя сокращать сценарий до обзора. "
        "Нужен подробный прикладной документ, по которому ведущий и DJ реально смогут работать на площадке. "
        "Работай как премиальный креативный помощник, а не как конвертер анкеты в блоки. "
        "Сначала делай выводы и режиссерские решения, потом пиши итог. "
        "Ищи уникальную зацепку пары и используй технику метафорического переноса. "
        "Если в истории есть лес, грибы, друг, музыка, длинный срок отношений или другая сильная деталь, превращай это в центральную метафору, а не в сухой факт. "
        "В каждом host_text должно быть не меньше 10 свежих, небанальных метафор. "
        "Каждый host_text и ключевые поля host_script пиши как готовую прямую речь ведущего от первого лица, а не как рекомендации о том, что ему делать. "
        "DJ guidance должен содержать не только трековые ориентиры, но и полноценный DJ sheet по точкам вечера: welcome, opening, family exit, first lift, dance block 1, dance block 2, late peak, final close. "
        "Нельзя писать общие слова вроде 'что-то атмосферное' или 'поднять зал'; нужны конкретные артистические и трековые ориентиры. "
        "Если в анкете указаны любимые песни или музыкальные предпочтения, обязательно встрои их в DJ-план как конкретные точки сета. "
        "Предложи несколько трендовых фишек ведущего, которые реально можно использовать на площадке без кринжа. "
        "Нельзя выдавать взаимозаменяемый сценарий. Каждая ключевая речь должна быть привязана к реальным деталям анкеты: истории знакомства, предложению, ценностям пары, их личному языку, важным датам, пожеланиям и фактуре гостей. "
        "Минимум в 4 ключевых блоках должны прозвучать конкретные уникальные детали пары, которые невозможно безболезненно переставить в чужую свадьбу. "
        "Запрещено повторять одинаковые композиции фраз между разными анкетами. Если абзац можно вставить в другую заявку почти без изменений, его нужно переписать глубже и конкретнее. "
        "Используй формулу detail -> meaning -> stage delivery. "
        "Не пересказывай анкету: превращай деталь в смысл, смысл в образ, образ в готовую живую реплику. "
        "Запрещены стоп-слова из negative_prompt_lexicon. "
        "Для ключевых вступительных и переходных текстов мысленно проверь 3 режима тона: эмоциональный, драйвовый и метафоричный. В итоговый JSON включай лучший вариант, но пиши с этой внутренней вариативностью. "
        "Нам не нужен шаблон и не нужна настройка. Нам нужен персональный помощник по этой анкете. "
        "Учитывай дату события и year-aware trend-bank: если событие в 2026 году, логика должна быть актуальной для 2026, а не для 2025."
    )


def build_compact_generation_user_prompt(questionnaire: dict[str, Any]) -> str:
    profile = build_personalization_brief(questionnaire)
    trend_bank = build_trend_bank(questionnaire)
    style_bank = build_style_bank(questionnaire)
    dramaturgy = build_dramaturgy_bank(questionnaire)
    return (
        f"Тип мероприятия: {questionnaire['eventType']}\n\n"
        "Анкета:\n"
        f"{build_questionnaire_context(questionnaire)}\n\n"
        "Краткое креативное ядро:\n"
        f"- Герои: {profile['heroes']}\n"
        f"- История: {profile['love_story']}\n"
        f"- Предложение / ключевой поворот: {profile['proposal_story']}\n"
        f"- Ценности: {', '.join(profile['values'][:4])}\n"
        f"- Внутренний язык: {profile['nicknames'] or 'использовать уместные личные обращения'}\n"
        f"- Внутренние шутки: {profile['inside_jokes'] or 'не выдумывать шутки без фактуры'}\n"
        f"- Важные даты: {', '.join(profile['important_dates'][:4]) or 'встроить реальные точки истории пары'}\n"
        f"- Любимая музыка: {', '.join(profile['music_likes'][:6]) or 'использовать музыкальные предпочтения клиента'}\n\n"
        "Dramaturgy-bank:\n"
        f"- Центральная метафора: {dramaturgy['central_metaphor']}\n"
        f"- Образный словарь: {'; '.join(dramaturgy['image_dictionary'])}\n"
        f"- Уникальные зацепки: {'; '.join([item for item in dramaturgy['unique_hooks'] if item])}\n"
        f"- Формула письма: {'; '.join(dramaturgy['writing_formula'])}\n"
        f"- Внутренние режимы тона: {'; '.join(dramaturgy['opening_variants'])}\n\n"
        "Trend-bank:\n"
        f"- Год события: {trend_bank['event_year']}\n"
        f"- Сезон: {trend_bank['season']}\n"
        f"- Сезонная логика: {trend_bank['seasonal_logic']}\n"
        f"- Host direction: {'; '.join(trend_bank['host_market_direction'])}\n"
        f"- DJ direction: {'; '.join(trend_bank['dj_market_direction'])}\n\n"
        "Style-bank:\n"
        f"- Voice direction: {style_bank['voice_direction']}\n"
        f"- Must sound like: {'; '.join(style_bank['must_sound_like'])}\n"
        f"- Must not sound like: {'; '.join(style_bank['must_not_sound_like'])}\n\n"
        "Верни только валидный JSON. "
        "Нужен премиальный сценарий, а не пересказ анкеты. "
        "Пиши как креативный помощник и режиссер, а не как шаблонный генератор. "
        "Собери реальный DJ sheet по точкам вечера и обязательно встрои любимые треки клиента в конкретные места сета. "
        "Работай через центральную метафору пары, detail -> meaning -> stage delivery и жестко избегай канцелярита и стоп-слов."
    )


def build_dossier_system_prompt() -> str:
    return """
Ты — сценарист-драматург и креативный стратег event-проекта.
Твоя задача: не писать финальный сценарий, а выделить драматургическое ядро пары и вечера.

Верни только JSON со структурой:
- central_metaphor
- image_dictionary
- unique_hooks
- forbidden_lexicon
- sacred_tracks
- dj_arc
- host_tone
- creative_risks
- writing_thesis

Правила:
- никакого канцелярита
- не пересказывай анкету подряд
- ищи 1 центральную метафору и 2-4 поддерживающих образа
- используй detail -> meaning -> stage delivery
- если есть музыкальные предпочтения, вытащи sacred tracks
"""


def build_dossier_user_prompt(questionnaire: dict[str, Any]) -> str:
    profile = json.dumps(build_personalization_brief(questionnaire), ensure_ascii=False, indent=2)
    trend_bank = json.dumps(build_trend_bank(questionnaire), ensure_ascii=False, indent=2)
    style_bank = json.dumps(build_style_bank(questionnaire), ensure_ascii=False, indent=2)
    dramaturgy_bank = json.dumps(build_dramaturgy_bank(questionnaire), ensure_ascii=False, indent=2)
    return (
        "Анкета:\n"
        f"{build_questionnaire_context(questionnaire)}\n\n"
        "Personalization brief:\n"
        f"{profile}\n\n"
        "Trend-bank:\n"
        f"{trend_bank}\n\n"
        "Style-bank:\n"
        f"{style_bank}\n\n"
        "Dramaturgy-bank:\n"
        f"{dramaturgy_bank}\n\n"
        "Собери creative dossier для этой пары и этого вечера. Верни только JSON."
    )


def build_writer_user_prompt(questionnaire: dict[str, Any], dossier: dict[str, Any]) -> str:
    return (
        f"{build_generation_user_prompt(questionnaire)}\n\n"
        "Creative dossier:\n"
        f"{json.dumps(dossier, ensure_ascii=False, indent=2)}\n\n"
        "Теперь напиши финальный рабочий JSON-сценарий, используя dossier как обязательное драматургическое ядро."
    )


WRITER_CHUNK_SPECS = [
    ("passport_quality", 42, "Собираем паспорт события и quality panel", "Верни только JSON с ключами event_passport и quality_panel."),
    ("concept_trend", 46, "Собираем концепцию и trend layer", "Верни только JSON с ключами concept и trend_layer."),
    ("commands_logic", 50, "Собираем команды ведущему, уточнения и режиссерскую логику", "Верни только JSON с ключами key_host_commands, questions_to_clarify_before_event и director_logic."),
    ("timeline_opening", 56, "Пишем первые 4 блока сценария", "Верни только JSON с ключом scenario_timeline. Внутри должны быть ровно 4 блока: открытие, первый смысловой разгон, вовлечение, подводка к развитию."),
    ("timeline_closing", 62, "Пишем вторые 4 блока сценария", "Верни только JSON с ключом scenario_timeline. Внутри должны быть ровно 4 блока: эмоциональное ядро, музыкальный подъем, поздний пик, финал."),
    ("host_script", 68, "Пишем ключевые тексты ведущего", "Верни только JSON с ключом host_script. Все поля должны быть прямой речью ведущего от первого лица."),
    ("dj_guidance", 73, "Собираем DJ sheet по предпочтениям пары", "Верни только JSON с ключом dj_guidance. Встрой sacred tracks и реальные музыкальные точки вечера."),
    ("supporting_sections", 78, "Собираем гостей, риски, план Б и короткую версию", "Верни только JSON с ключами guest_management, risk_map, plan_b и final_print_version."),
]

DOSSIER_CHUNK_SPECS = [
    ("story_dna", 12, "Собираем драматургическое ядро пары", "Верни только JSON с ключами central_metaphor, supporting_images, unique_hook, couple_dynamics, tone_profile, forbidden_cliches, personalization_mandates."),
    ("music_guest_dna", 18, "Собираем музыкальную и гостевую ДНК вечера", "Верни только JSON с ключами sacred_tracks, music_map, guest_energy_map, sensitive_zones, host_trend_features, staging_priorities."),
]

POLISH_CHUNK_SPECS = [
    ("polish_timeline", 82, "Шлифуем таймлайн и прямую речь ведущего", "Верни только JSON с ключами scenario_timeline и host_script. Усиль персонализацию и убери пересказ анкеты."),
    ("polish_dj", 88, "Шлифуем DJ sheet и музыкальные переходы", "Верни только JSON с ключом dj_guidance. Сделай материал конкретным по трекам, переходам и rescue-логике."),
    ("polish_support", 94, "Шлифуем риски, план Б и печатную версию", "Верни только JSON с ключами guest_management, risk_map, plan_b и final_print_version."),
]


def build_writer_chunk_user_prompt(
    questionnaire: dict[str, Any],
    dossier: dict[str, Any],
    partial_program: dict[str, Any],
    chunk_name: str,
    chunk_instruction: str,
) -> str:
    already_written = json.dumps(partial_program, ensure_ascii=False, indent=2) if partial_program else "{}"
    return (
        "Анкета:\n"
        f"{build_questionnaire_context(questionnaire)}\n\n"
        "Creative dossier:\n"
        f"{json.dumps(dossier, ensure_ascii=False, indent=2)}\n\n"
        "Уже написанные части сценария:\n"
        f"{already_written}\n\n"
        f"Текущий writer-chunk: {chunk_name}\n"
        f"{chunk_instruction}\n\n"
        "Правила chunk-генерации:\n"
        "- не повторяй композиции, одинаковые первые фразы и одну и ту же метафору между блоками\n"
        "- превращай деталь анкеты в сценическое решение, а не в пересказ\n"
        "- если пишешь timeline, каждый блок должен отличаться по задаче, ритму и языку\n"
        "- если пишешь host_script, он должен усиливать timeline, а не копировать его\n"
        "- если пишешь dj_guidance, используй любимые треки пары, sacred tracks и конкретные музыкальные точки\n"
        "- верни только валидный JSON для текущего chunk"
    )


def build_dossier_chunk_user_prompt(
    questionnaire: dict[str, Any],
    partial_dossier: dict[str, Any],
    chunk_name: str,
    chunk_instruction: str,
) -> str:
    return (
        "Анкета:\n"
        f"{build_questionnaire_context(questionnaire)}\n\n"
        "Уже собранный dossier:\n"
        f"{json.dumps(partial_dossier, ensure_ascii=False, indent=2)}\n\n"
        f"Текущий dossier-chunk: {chunk_name}\n"
        f"{chunk_instruction}\n\n"
        "Правила:\n"
        "- не пересказывай анкету; делай выводы\n"
        "- ищи центральную метафору, драматургию и прикладные выводы\n"
        "- верни только валидный JSON для текущего chunk"
    )


def build_polish_chunk_user_prompt(
    questionnaire: dict[str, Any],
    dossier: dict[str, Any],
    program: dict[str, Any],
    chunk_name: str,
    chunk_instruction: str,
) -> str:
    return (
        "Анкета:\n"
        f"{build_questionnaire_context(questionnaire)}\n\n"
        "Creative dossier:\n"
        f"{json.dumps(dossier, ensure_ascii=False, indent=2)}\n\n"
        "Текущий сценарий:\n"
        f"{json.dumps(program, ensure_ascii=False, indent=2)}\n\n"
        f"Текущий polish-chunk: {chunk_name}\n"
        f"{chunk_instruction}\n\n"
        "Правила:\n"
        "- меняй только нужные секции\n"
        "- не переписывай весь документ целиком\n"
        "- усиливай персонализацию и конкретику\n"
        "- верни только валидный JSON для текущего chunk"
    )


def merge_program_chunk(program: dict[str, Any], chunk: dict[str, Any]) -> dict[str, Any]:
    merged = dict(program)
    for key, value in chunk.items():
        if key == "scenario_timeline" and isinstance(value, list):
            existing = as_list(merged.get("scenario_timeline"))
            existing.extend(value)
            merged["scenario_timeline"] = existing
        elif isinstance(value, dict) and isinstance(merged.get(key), dict):
            current = as_dict(merged.get(key))
            current.update(value)
            merged[key] = current
        else:
            merged[key] = value
    return merged


def build_polish_system_prompt() -> str:
    return """
Ты — финальный редактор и script doctor.
Твоя задача: не переписывать все с нуля, а улучшить уже собранный сценарий.

Верни только JSON той же структуры.

Что улучшить:
- убрать канцелярит и шаблонные куски
- усилить личные детали пары
- проверить, что DJ guidance реально привязан к предпочтениям клиента
- сделать тексты ведущего более пригодными к живому произнесению
- сохранить структуру и прикладность документа
"""


def build_polish_user_prompt(questionnaire: dict[str, Any], dossier: dict[str, Any], program: dict[str, Any]) -> str:
    return (
        "Анкета:\n"
        f"{build_questionnaire_context(questionnaire)}\n\n"
        "Creative dossier:\n"
        f"{json.dumps(dossier, ensure_ascii=False, indent=2)}\n\n"
        "Текущий сценарий:\n"
        f"{json.dumps(program, ensure_ascii=False, indent=2)}\n\n"
        "Отредактируй сценарий как финальный script doctor и верни только улучшенный JSON."
    )


def normalize_program(program: dict[str, Any], questionnaire: dict[str, Any]) -> dict[str, Any]:
    fallback = {
        "event_passport": {
            "event_type": questionnaire.get("eventType", ""),
            "event_date": questionnaire.get("eventDate", ""),
            "city": questionnaire.get("city", ""),
            "venue": questionnaire.get("venue", ""),
            "timing_anchor": questionnaire.get("startTime", ""),
        },
        "quality_panel": {},
        "concept": {},
        "trend_layer": {},
        "key_host_commands": [],
        "questions_to_clarify_before_event": [],
        "director_logic": {},
        "scenario_timeline": [],
        "host_script": {},
        "dj_guidance": {},
        "guest_management": {},
        "risk_map": [],
        "plan_b": [],
        "final_print_version": {},
    }
    for key, value in fallback.items():
        if key not in program or not program[key]:
            program[key] = value
    for key in [
        "event_passport",
        "quality_panel",
        "concept",
        "trend_layer",
        "director_logic",
        "host_script",
        "dj_guidance",
        "guest_management",
        "final_print_version",
    ]:
        if not isinstance(program.get(key), dict):
            program[key] = fallback[key]
    for key in ["key_host_commands", "questions_to_clarify_before_event", "risk_map", "plan_b"]:
        if not isinstance(program.get(key), list):
            program[key] = fallback[key]
    if not isinstance(program.get("scenario_timeline"), list) or not program["scenario_timeline"]:
        timing_anchor = questionnaire.get("startTime", "").strip() or "18:00"
        program["scenario_timeline"] = [
            {
                "time_from": timing_anchor,
                "time_to": timing_anchor,
                "block_title": "Основа вечера",
                "block_purpose": "Собрать рабочий каркас сценария для ведущего.",
                "what_happens": "Ведущий открывает вечер и задает логику дальнейшего развития.",
                "host_action": "Держать тон, внимание зала и мягко вести к следующим блокам.",
                "host_text": "Добрый вечер. Сегодня мы собираем не формальный набор блоков, а живую историю этого события.",
                "dj_task": "Держать нейтральную аккуратную подложку без резких входов.",
                "director_move": "Сначала собрать внимание, затем расширять драматургию.",
                "risk_control": "Не торопить зал и не перегружать старт объявлениями.",
                "transition": "После открытия перейти к первому смысловому блоку.",
            }
        ]
    return program


def finalize_generated_program(
    program: Any,
    questionnaire: dict[str, Any],
    dossier: dict[str, Any] | None = None,
) -> dict[str, Any]:
    repaired = normalize_program(as_dict(program), questionnaire)
    repaired = refine_program_payload(repaired, questionnaire)
    repaired = normalize_program(repaired, questionnaire)
    repaired["_schema_version"] = PROGRAM_SCHEMA_VERSION
    if dossier is not None:
        repaired["_creative_dossier"] = dossier

    if is_program_actual(repaired):
        return repaired

    fallback = normalize_program({}, questionnaire)
    fallback["_schema_version"] = PROGRAM_SCHEMA_VERSION
    if dossier is not None:
        fallback["_creative_dossier"] = dossier
    return fallback


def annotate_program_source(program: dict[str, Any], *, source: str, note: str = "") -> dict[str, Any]:
    program["_generation_meta"] = {
        "source": source,
        "mode": EVENT_AI_MODE,
        "model": AI_MODEL,
        "reasoning_effort": AI_REASONING_EFFORT,
        "strict_ai_only": STRICT_AI_ONLY,
        "note": note,
        "generated_at": datetime.now().isoformat(),
    }
    return program


def annotate_dossier(dossier: dict[str, Any]) -> dict[str, Any]:
    dossier["_meta"] = {
        "mode": EVENT_AI_MODE,
        "model": AI_MODEL,
        "generated_at": datetime.now().isoformat(),
    }
    return dossier


def is_timeout_error(error: Exception) -> bool:
    text = str(error).lower()
    return "timed out" in text or "timeout" in text or "time out" in text


def is_gateway_error(error: Exception) -> bool:
    text = str(error).lower()
    return (
        "502 bad gateway" in text
        or "503 service unavailable" in text
        or "504 gateway timeout" in text
        or "bad gateway" in text
    )


def request_program_from_openai(*, system_prompt: str, user_prompt: str, model: str, reasoning_effort: str) -> dict[str, Any]:
    if client is None:
        raise RuntimeError("OpenAI client unavailable")
    last_error: Exception | None = None
    attempts = max(1, OPENAI_GATEWAY_RETRIES)
    for attempt in range(1, attempts + 1):
        try:
            response = client.responses.create(
                model=model,
                reasoning={"effort": reasoning_effort},
                input=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
            )
            raw_text = response.output_text.strip()
            return parse_json_response(raw_text)
        except Exception as error:
            last_error = error
            if not is_gateway_error(error) or attempt >= attempts:
                break
            sleep(OPENAI_GATEWAY_RETRY_DELAY_SECONDS * attempt)
    if last_error is not None:
        raise RuntimeError(f"OpenAI upstream error after retries: {last_error}")
    raise RuntimeError("OpenAI request failed without explicit error")


def generate_creative_dossier(
    questionnaire: dict[str, Any],
    progress_callback: Any | None = None,
) -> dict[str, Any]:
    dossier: dict[str, Any] = {}
    for chunk_name, percent, status_message, chunk_instruction in DOSSIER_CHUNK_SPECS:
        if progress_callback is not None:
            progress_callback(percent, status_message, dossier)
        chunk = request_program_from_openai(
            system_prompt=build_dossier_system_prompt(),
            user_prompt=build_dossier_chunk_user_prompt(questionnaire, dossier, chunk_name, chunk_instruction),
            model=OPENAI_DOSSIER_MODEL,
            reasoning_effort=OPENAI_DOSSIER_REASONING,
        )
        dossier.update(as_dict(chunk))
    return annotate_dossier(dossier)


def write_program_from_dossier(
    questionnaire: dict[str, Any],
    dossier: dict[str, Any],
    progress_callback: Any | None = None,
) -> dict[str, Any]:
    program: dict[str, Any] = {}
    for chunk_name, percent, status_message, chunk_instruction in WRITER_CHUNK_SPECS:
        if progress_callback is not None:
            progress_callback(percent, status_message, program)
        chunk = request_program_from_openai(
            system_prompt=build_stage_system_prompt(),
            user_prompt=build_writer_chunk_user_prompt(questionnaire, dossier, program, chunk_name, chunk_instruction),
            model=OPENAI_WRITER_MODEL,
            reasoning_effort=OPENAI_WRITER_REASONING,
        )
        program = merge_program_chunk(program, as_dict(chunk))
        if progress_callback is not None:
            progress_callback(percent, status_message, program)
    return normalize_program(program, questionnaire)


def polish_program(
    questionnaire: dict[str, Any],
    dossier: dict[str, Any],
    program: dict[str, Any],
    progress_callback: Any | None = None,
) -> dict[str, Any]:
    polished = dict(program)
    for chunk_name, percent, status_message, chunk_instruction in POLISH_CHUNK_SPECS:
        if progress_callback is not None:
            progress_callback(percent, status_message, polished)
        chunk = request_program_from_openai(
            system_prompt=build_polish_system_prompt(),
            user_prompt=build_polish_chunk_user_prompt(questionnaire, dossier, polished, chunk_name, chunk_instruction),
            model=OPENAI_POLISH_MODEL,
            reasoning_effort=OPENAI_POLISH_REASONING,
        )
        polished = merge_program_chunk(polished, as_dict(chunk))
        if progress_callback is not None:
            progress_callback(percent, status_message, polished)
    return normalize_program(polished, questionnaire)


def parse_json_response(raw_text: str) -> dict[str, Any]:
    try:
        return json.loads(raw_text)
    except json.JSONDecodeError:
        start = raw_text.find("{")
        end = raw_text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise RuntimeError("OpenAI вернул невалидный JSON")
        return json.loads(raw_text[start : end + 1])


def is_program_actual(program: Any) -> bool:
    if not isinstance(program, dict):
        return False
    if program.get("_schema_version") != PROGRAM_SCHEMA_VERSION:
        return False
    required_keys = {
        "event_passport",
        "quality_panel",
        "concept",
        "trend_layer",
        "key_host_commands",
        "questions_to_clarify_before_event",
        "director_logic",
        "scenario_timeline",
        "host_script",
        "dj_guidance",
        "guest_management",
        "risk_map",
        "plan_b",
        "final_print_version",
    }
    if not required_keys.issubset(program.keys()):
        return False
    if not isinstance(program.get("scenario_timeline"), list) or len(program.get("scenario_timeline", [])) < 1:
        return False
    return True


def generate_agent_program(questionnaire: dict[str, Any]) -> dict[str, Any]:
    fallback_program = build_target_program(questionnaire)
    fallback_program["_schema_version"] = PROGRAM_SCHEMA_VERSION
    if client is None:
        if STRICT_AI_ONLY:
            raise RuntimeError("OPENAI_API_KEY не настроен или OpenAI клиент недоступен. Premium-режим не будет подменять результат шаблонным fallback.")
        return annotate_program_source(
            fallback_program,
            source="fallback",
            note="OpenAI client unavailable; returned fallback program.",
        )

    try:
        program = request_program_from_openai(
            system_prompt=build_stage_system_prompt(),
            user_prompt=(
                f"Тип мероприятия: {questionnaire['eventType']}\n\n"
                "Анкета:\n"
                f"{build_questionnaire_context(questionnaire)}\n\n"
                "Верни только валидный JSON."
            ),
            model=AI_MODEL,
            reasoning_effort=AI_REASONING_EFFORT,
        )
        program = normalize_program(program, questionnaire)
        program["_schema_version"] = PROGRAM_SCHEMA_VERSION
        return annotate_program_source(program, source="openai", note="Program generated through OpenAI Responses API.")
    except Exception as error:
        if is_timeout_error(error) and OPENAI_TIMEOUT_RETRY_MODEL:
            try:
                program = request_program_from_openai(
                    system_prompt=build_stage_system_prompt(),
                    user_prompt=build_compact_generation_user_prompt(questionnaire),
                    model=OPENAI_TIMEOUT_RETRY_MODEL,
                    reasoning_effort=OPENAI_TIMEOUT_RETRY_REASONING,
                )
                program = normalize_program(program, questionnaire)
                program["_schema_version"] = PROGRAM_SCHEMA_VERSION
                return annotate_program_source(
                    program,
                    source="openai-retry",
                    note=f"Primary request timed out; retry succeeded with {OPENAI_TIMEOUT_RETRY_MODEL}.",
                )
            except Exception as retry_error:
                error = retry_error
        if STRICT_AI_ONLY:
            raise RuntimeError(f"OpenAI генерация не сработала: {error}")
        return annotate_program_source(
            fallback_program,
            source="fallback",
            note=f"OpenAI generation failed; fallback returned: {error}",
        )


def generate_agent_program_fast(questionnaire: dict[str, Any]) -> dict[str, Any]:
    fallback_program = build_target_program(questionnaire)
    fallback_program["_schema_version"] = PROGRAM_SCHEMA_VERSION
    if client is None:
        if STRICT_AI_ONLY:
            raise RuntimeError("OPENAI_API_KEY не настроен или OpenAI клиент недоступен. Premium-режим не будет подменять результат шаблонным fallback.")
        return annotate_program_source(
            fallback_program,
            source="fallback",
            note="OpenAI client unavailable; returned fallback program.",
        )

    try:
        program = request_program_from_openai(
            system_prompt=build_stage_system_prompt(),
            user_prompt=build_generation_user_prompt(questionnaire),
            model=AI_MODEL,
            reasoning_effort=AI_REASONING_EFFORT,
        )
        program = normalize_program(program, questionnaire)
        program["_schema_version"] = PROGRAM_SCHEMA_VERSION
        return annotate_program_source(program, source="openai", note="Program generated through OpenAI Responses API.")
    except Exception as error:
        if is_timeout_error(error) and OPENAI_TIMEOUT_RETRY_MODEL:
            try:
                program = request_program_from_openai(
                    system_prompt=build_stage_system_prompt(),
                    user_prompt=build_compact_generation_user_prompt(questionnaire),
                    model=OPENAI_TIMEOUT_RETRY_MODEL,
                    reasoning_effort=OPENAI_TIMEOUT_RETRY_REASONING,
                )
                program = normalize_program(program, questionnaire)
                program["_schema_version"] = PROGRAM_SCHEMA_VERSION
                return annotate_program_source(
                    program,
                    source="openai-retry",
                    note=f"Primary request timed out; retry succeeded with {OPENAI_TIMEOUT_RETRY_MODEL}.",
                )
            except Exception as retry_error:
                error = retry_error
        if STRICT_AI_ONLY:
            raise RuntimeError(f"OpenAI генерация не сработала: {error}")
        return annotate_program_source(
            fallback_program,
            source="fallback",
            note=f"OpenAI generation failed; fallback returned: {error}",
        )


def safe_filename(value: str) -> str:
    translit_map = {
        "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "e",
        "ж": "zh", "з": "z", "и": "i", "й": "y", "к": "k", "л": "l", "м": "m",
        "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t", "у": "u",
        "ф": "f", "х": "h", "ц": "ts", "ч": "ch", "ш": "sh", "щ": "sch", "ы": "y",
        "э": "e", "ю": "yu", "я": "ya", "ь": "", "ъ": "",
    }
    value = value.strip().lower()
    result: list[str] = []
    for char in value:
        if char in translit_map:
            result.append(translit_map[char])
        elif char.isalnum():
            result.append(char)
        elif char in {" ", "-", "_"}:
            result.append("_")
    return re.sub(r"_+", "_", "".join(result)).strip("_") or "event_ai"


def add_label_value(document: Document, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(str(value).strip() if str(value).strip() else "Не указано")


def add_list(document: Document, items: list[Any]) -> None:
    if not items:
        document.add_paragraph("Не указано")
        return
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def build_docx(submission: dict[str, Any], program: dict[str, Any]) -> BytesIO:
    questionnaire = submission["questionnaire"]
    labels = get_questionnaire_labels()

    document = Document()
    document.add_heading("Event AI: анкета и программа ведущего", 0)
    add_label_value(document, "ID заявки", submission["id"])
    add_label_value(document, "Создано", submission["created_at"])

    document.add_heading("1. Анкета клиента", level=1)
    for key, label in labels.items():
        add_label_value(document, label, questionnaire.get(key, ""))

    document.add_heading("2. Сценарий ведущего", level=1)

    passport = as_dict(program.get("event_passport"))
    document.add_heading("2.1 Паспорт события", level=2)
    for field, title in [
        ("event_type", "Тип"),
        ("format_name", "Формат"),
        ("city", "Город"),
        ("venue", "Площадка"),
        ("event_date", "Дата"),
        ("working_timeline_note", "Рабочая пометка по таймингу"),
        ("main_goal", "Главная цель"),
        ("atmosphere", "Атмосфера"),
        ("style", "Стиль"),
        ("timing_anchor", "Тайминговый якорь"),
    ]:
        add_label_value(document, title, passport.get(field, ""))
    document.add_paragraph("Обязательные точки:")
    add_list(document, as_list(passport.get("mandatory_points")))
    document.add_paragraph("Жесткие запреты:")
    add_list(document, as_list(passport.get("hard_bans")))

    quality = as_dict(program.get("quality_panel"))
    document.add_heading("2.2 Качество результата", level=2)
    add_label_value(document, "Сценарный вердикт", quality.get("scenario_verdict", ""))
    add_label_value(document, "Режиссерский вердикт", quality.get("director_verdict", ""))
    add_label_value(document, "Критический вердикт", quality.get("critic_verdict", ""))
    add_label_value(document, "Готово к работе", "Да" if quality.get("final_ready") else "Нет")
    document.add_paragraph("Исправленные слабые места:")
    add_list(document, as_list(quality.get("fixed_issues")))

    concept = as_dict(program.get("concept"))
    document.add_heading("2.3 Концепция", level=2)
    for field, title in [
        ("big_idea", "Большая идея"),
        ("main_director_thesis", "Главный режиссерский тезис"),
        ("main_emotional_result", "Эмоциональный результат"),
        ("why_this_event_will_be_remembered", "Почему вечер запомнится"),
    ]:
        add_label_value(document, title, concept.get(field, ""))

    trend_layer = as_dict(program.get("trend_layer"))
    document.add_heading("2.4 Современная логика", level=2)
    add_label_value(document, "Резюме", trend_layer.get("trend_summary", ""))
    document.add_paragraph("Что применено:")
    add_list(document, as_list(trend_layer.get("applied_trends")))
    document.add_paragraph("Что отброшено как устаревшее:")
    add_list(document, as_list(trend_layer.get("rejected_outdated_patterns")))

    document.add_heading("2.5 Ключевые команды ведущему", level=2)
    add_list(document, as_list(program.get("key_host_commands")))

    document.add_heading("2.6 Что нужно уточнить до мероприятия", level=2)
    add_list(document, as_list(program.get("questions_to_clarify_before_event")))

    logic = as_dict(program.get("director_logic"))
    document.add_heading("2.7 Режиссерская логика вечера", level=2)
    for field, title in [
        ("opening_logic", "Логика открытия"),
        ("development_logic", "Логика развития"),
        ("family_or_core_emotional_logic", "Логика эмоционального ядра"),
        ("final_logic", "Логика финала"),
    ]:
        add_label_value(document, title, logic.get(field, ""))

    document.add_heading("2.8 Пошаговый тайминг", level=2)
    for index, block in enumerate(as_list(program.get("scenario_timeline")), start=1):
        block = as_dict(block)
        document.add_paragraph(f"{index}. {block.get('time_from', '')} - {block.get('time_to', '')}: {block.get('block_title', '')}")
        for field, title in [
            ("block_purpose", "Цель"),
            ("what_happens", "Что происходит"),
            ("host_action", "Действие ведущего"),
            ("host_text", "Текст ведущего"),
            ("dj_task", "Задача диджея"),
            ("director_move", "Режиссерский ход"),
            ("risk_control", "Контроль риска"),
            ("transition", "Переход"),
        ]:
            add_label_value(document, title, block.get(field, ""))

    script = as_dict(program.get("host_script"))
    document.add_heading("2.9 Тексты ведущего", level=2)
    for field, title in [
        ("opening_main", "Opening main"),
        ("opening_short", "Opening short"),
        ("welcome_line", "Welcome line"),
        ("first_core_intro", "First core intro"),
        ("family_block_intro", "Family block intro"),
        ("surprise_intro", "Surprise intro"),
        ("dance_block_intro", "Dance block intro"),
        ("final_block_intro", "Final block intro"),
        ("closing_words", "Closing words"),
    ]:
        add_label_value(document, title, script.get(field, ""))

    dj = as_dict(program.get("dj_guidance"))
    document.add_heading("2.10 DJ guidance", level=2)
    for field, title in [
        ("overall_music_policy", "Overall music policy"),
        ("welcome_music", "Welcome music"),
        ("opening_music", "Opening music"),
        ("table_background", "Table background"),
        ("emotional_blocks_music", "Emotional blocks music"),
        ("dance_block_1", "Dance block 1"),
        ("dance_block_2", "Dance block 2"),
        ("dance_block_3", "Dance block 3"),
        ("final_block_music", "Final block music"),
        ("final_music", "Final music"),
    ]:
        add_label_value(document, title, dj.get(field, ""))
    document.add_paragraph("Stop list:")
    add_list(document, as_list(dj.get("stop_list")))
    document.add_paragraph("Technical notes:")
    add_list(document, as_list(dj.get("technical_notes")))

    guest_management = as_dict(program.get("guest_management"))
    document.add_heading("2.11 Работа с гостями", level=2)
    for field, title in [
        ("active_people", "Активные люди"),
        ("shy_people", "Скромные люди"),
        ("important_people", "Важные люди"),
        ("do_not_involve", "Кого не вовлекать"),
        ("sensitive_people_or_topics", "Чувствительные люди и темы"),
        ("management_notes", "Рабочие заметки"),
    ]:
        document.add_paragraph(f"{title}:")
        add_list(document, as_list(guest_management.get(field)))

    document.add_heading("2.12 Риски", level=2)
    for item in as_list(program.get("risk_map")):
        item = as_dict(item)
        add_label_value(document, "Риск", item.get("risk", ""))
        add_label_value(document, "Почему важен", item.get("why_it_matters", ""))
        add_label_value(document, "Как предотвратить", item.get("how_to_prevent", ""))
        add_label_value(document, "Что делать, если сработало", item.get("what_to_do_if_triggered", ""))

    document.add_heading("2.13 План Б", level=2)
    for item in as_list(program.get("plan_b")):
        item = as_dict(item)
        add_label_value(document, "Ситуация", item.get("situation", ""))
        add_label_value(document, "Решение", item.get("solution", ""))

    final_print = as_dict(program.get("final_print_version"))
    document.add_heading("2.14 Краткая версия для печати", level=2)
    add_label_value(document, "Название", final_print.get("title", ""))
    add_label_value(document, "Краткое резюме", final_print.get("summary", ""))
    for field, title in [
        ("timeline_short", "Короткий таймлайн"),
        ("must_do", "Must do"),
        ("must_not_do", "Must not do"),
        ("host_focus", "Фокус ведущего"),
        ("dj_focus", "Фокус диджея"),
    ]:
        document.add_paragraph(f"{title}:")
        add_list(document, as_list(final_print.get(field)))

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def default_generation_state() -> dict[str, Any]:
    return {
        "status": "idle",
        "stage": "",
        "percent": 0,
        "message": "",
        "error": "",
        "job_id": "",
        "updated_at": datetime.now().isoformat(),
    }


def save_generation_state(
    submission_id: str,
    *,
    status: str,
    stage: str = "",
    percent: int | None = None,
    message: str = "",
    error: str = "",
    program: dict[str, Any] | None = None,
    job_id: str | None = None,
    expected_job_id: str | None = None,
) -> None:
    submissions = load_submissions()
    for index, submission in enumerate(submissions):
        if submission["id"] != submission_id:
            continue
        generation = submission.get("generation")
        if not isinstance(generation, dict):
            generation = default_generation_state()
        current_job_id = str(generation.get("job_id", "") or "")
        if expected_job_id is not None and current_job_id != expected_job_id:
            return
        generation.update(
            {
                "status": status,
                "stage": stage,
                "message": message,
                "error": error,
                "updated_at": datetime.now().isoformat(),
            }
        )
        if percent is not None:
            generation["percent"] = max(0, min(100, int(percent)))
        if job_id is not None:
            generation["job_id"] = job_id
        submissions[index]["generation"] = generation
        if program is not None:
            submissions[index]["program"] = program
        save_submissions(submissions)
        return


def run_generation_job(submission_id: str, job_id: str) -> None:
    try:
        submissions, index = get_submission_or_404(submission_id)
        questionnaire = submissions[index]["questionnaire"]
        save_generation_state(
            submission_id,
            status="running",
            stage="dossier",
            percent=10,
            message="Ищем золотые зацепки пары, центральную метафору и sacred tracks",
            expected_job_id=job_id,
        )
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(
                generate_creative_dossier,
                questionnaire,
                lambda percent, message, _: save_generation_state(
                    submission_id,
                    status="running",
                    stage="dossier",
                    percent=percent,
                    message=message,
                    expected_job_id=job_id,
                ),
            )
            try:
                dossier = future.result(timeout=GENERATION_WATCHDOG_SECONDS)
            except FutureTimeoutError:
                future.cancel()
                save_generation_state(
                    submission_id,
                    status="failed",
                    stage="failed",
                    percent=100,
                    error=f"Генерация превысила лимит {int(GENERATION_WATCHDOG_SECONDS)} сек и была остановлена watchdog.",
                    message="Генерация заняла слишком много времени. Попробуйте еще раз.",
                    job_id=f"expired:{job_id}",
                    expected_job_id=job_id,
                )
                return
        save_generation_state(
            submission_id,
            status="running",
            stage="writer",
            percent=40,
            message="Пишем основной сценарий и музыкальную драматургию по creative dossier",
            expected_job_id=job_id,
        )
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(
                write_program_from_dossier,
                questionnaire,
                dossier,
                lambda percent, message, partial_program: save_generation_state(
                    submission_id,
                    status="running",
                    stage="writer",
                    percent=percent,
                    message=message,
                    program=finalize_generated_program(partial_program, questionnaire, dossier) if partial_program else None,
                    expected_job_id=job_id,
                ),
            )
            try:
                draft_program = future.result(timeout=GENERATION_WATCHDOG_SECONDS)
            except FutureTimeoutError:
                future.cancel()
                save_generation_state(
                    submission_id,
                    status="failed",
                    stage="failed",
                    percent=100,
                    error=f"Writer-этап превысил лимит {int(GENERATION_WATCHDOG_SECONDS)} сек и был остановлен watchdog.",
                    message="Основной writer-этап занял слишком много времени. Попробуйте еще раз.",
                    job_id=f"expired:{job_id}",
                    expected_job_id=job_id,
                )
                return
        save_generation_state(
            submission_id,
            status="running",
            stage="polish",
            percent=78,
            message="Убираем канцелярит, усиливаем образы и шлифуем DJ sheet",
            expected_job_id=job_id,
        )
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(
                polish_program,
                questionnaire,
                dossier,
                draft_program,
                lambda percent, message, partial_program: save_generation_state(
                    submission_id,
                    status="running",
                    stage="polish",
                    percent=percent,
                    message=message,
                    program=finalize_generated_program(partial_program, questionnaire, dossier),
                    expected_job_id=job_id,
                ),
            )
            try:
                program = future.result(timeout=GENERATION_WATCHDOG_SECONDS)
            except FutureTimeoutError:
                future.cancel()
                program = draft_program
        program = finalize_generated_program(program, questionnaire, dossier)
        save_generation_state(
            submission_id,
            status="ready",
            stage="final_assembly",
            percent=100,
            message="Программа готова",
            program=program,
            expected_job_id=job_id,
        )
    except Exception as error:
        save_generation_state(
            submission_id,
            status="failed",
            stage="failed",
            percent=100,
            error=str(error),
            message="Генерация завершилась с ошибкой",
            job_id=f"failed:{job_id}",
            expected_job_id=job_id,
        )
    finally:
        with GENERATION_THREAD_LOCK:
            GENERATION_THREADS.pop(submission_id, None)


def start_generation_job(submission_id: str) -> None:
    with GENERATION_THREAD_LOCK:
        existing = GENERATION_THREADS.get(submission_id)
        if existing and existing.is_alive():
            return
        job_id = uuid4().hex
        save_generation_state(
            submission_id,
            status="queued",
            stage="queued",
            percent=3,
            message="Задача поставлена в очередь генерации",
            error="",
            job_id=job_id,
        )
        worker = Thread(target=run_generation_job, args=(submission_id, job_id), daemon=True)
        GENERATION_THREADS[submission_id] = worker
        worker.start()


def get_submission_or_404(submission_id: str) -> tuple[list[dict[str, Any]], int]:
    submissions = load_submissions()
    for index, submission in enumerate(submissions):
        if submission["id"] == submission_id:
            return submissions, index
    raise HTTPException(status_code=404, detail="Анкета не найдена")


@app.get("/")
def root() -> dict[str, Any]:
    return {
        "status": "ok",
        "message": "Event AI backend is running",
        "data_path": str(SUBMISSIONS_FILE),
        "using_railway_volume": bool(RAILWAY_VOLUME_MOUNT_PATH),
        "mode": EVENT_AI_MODE,
        "model": AI_MODEL,
        "reasoning_effort": AI_REASONING_EFFORT,
        "strict_ai_only": STRICT_AI_ONLY,
        "openai_timeout_seconds": OPENAI_TIMEOUT_SECONDS,
        "openai_gateway_retries": OPENAI_GATEWAY_RETRIES,
        "generation_watchdog_seconds": GENERATION_WATCHDOG_SECONDS,
        "openai_enabled": bool(client),
        "supported_event_types": sorted(SUPPORTED_EVENT_TYPES),
    }


@app.get("/health")
def health() -> dict[str, Any]:
    return {
        "status": "healthy",
        "mode": EVENT_AI_MODE,
        "model": AI_MODEL,
        "reasoning_effort": AI_REASONING_EFFORT,
        "strict_ai_only": STRICT_AI_ONLY,
        "openai_timeout_seconds": OPENAI_TIMEOUT_SECONDS,
        "openai_gateway_retries": OPENAI_GATEWAY_RETRIES,
        "generation_watchdog_seconds": GENERATION_WATCHDOG_SECONDS,
        "openai_enabled": bool(client),
    }


@app.post("/api/questionnaire")
def submit_questionnaire(payload: QuestionnaireSubmission) -> dict[str, Any]:
    submission = {
        "id": datetime.now().strftime("%Y%m%d%H%M%S%f"),
        "created_at": datetime.now().isoformat(),
        "questionnaire": payload.model_dump(),
        "generation": default_generation_state(),
    }
    submissions = load_submissions()
    submissions.append(submission)
    save_submissions(submissions)
    return {
        "status": "success",
        "message": "Анкета сохранена",
        "clientName": payload.clientName,
        "eventType": payload.eventType,
        "savedId": submission["id"],
    }


@app.get("/api/submissions")
def get_submissions() -> dict[str, Any]:
    submissions = load_submissions()
    return {"status": "success", "count": len(submissions), "items": submissions}


@app.get("/api/submissions/{submission_id}")
def get_submission(submission_id: str) -> dict[str, Any]:
    submissions, index = get_submission_or_404(submission_id)
    return {"status": "success", "item": submissions[index]}


@app.get("/api/submissions/{submission_id}/generation-status")
def get_generation_status(submission_id: str) -> dict[str, Any]:
    submissions, index = get_submission_or_404(submission_id)
    submission = submissions[index]
    generation = submission.get("generation")
    if not isinstance(generation, dict):
        generation = default_generation_state()
    has_program = is_program_actual(submission.get("program"))
    if generation.get("status") == "ready" and not has_program:
        generation = {
            **generation,
            "status": "failed",
            "stage": "failed",
            "error": generation.get("error") or "Статус ready был выставлен без валидной программы.",
            "message": "Итоговая программа не прошла проверку. Запустите генерацию повторно.",
        }
    return {
        "status": "success",
        "submissionId": submission_id,
        "generation": generation,
        "hasProgram": has_program,
    }


@app.post("/api/submissions/{submission_id}/generate-program/start")
def start_generate_program(submission_id: str) -> dict[str, Any]:
    submissions, index = get_submission_or_404(submission_id)
    submission = submissions[index]
    if is_program_actual(submission.get("program")):
        save_generation_state(submission_id, status="ready", stage="final_assembly", percent=100, message="Программа уже готова")
        submissions, index = get_submission_or_404(submission_id)
        return {
            "status": "success",
            "submissionId": submission_id,
            "generation": submissions[index]["generation"],
            "cached": True,
            "program": submissions[index]["program"],
        }

    save_generation_state(submission_id, status="queued", stage="queued", percent=3, message="Ставим генерацию в работу")
    start_generation_job(submission_id)
    submissions, index = get_submission_or_404(submission_id)
    return {
        "status": "accepted",
        "submissionId": submission_id,
        "generation": submissions[index]["generation"],
        "cached": False,
    }


@app.post("/api/submissions/{submission_id}/generate-program")
def generate_program(submission_id: str) -> dict[str, Any]:
    submissions, index = get_submission_or_404(submission_id)
    submission = submissions[index]
    if is_program_actual(submission.get("program")):
        return {
            "status": "success",
            "submissionId": submission_id,
            "program": submission["program"],
            "cached": True,
        }

    try:
        save_generation_state(submission_id, status="running", stage="dossier", percent=10, message="Собираем creative dossier")
        dossier = generate_creative_dossier(
            submission["questionnaire"],
            lambda percent, message, _: save_generation_state(
                submission_id,
                status="running",
                stage="dossier",
                percent=percent,
                message=message,
            ),
        )
        save_generation_state(submission_id, status="running", stage="writer", percent=40, message="Пишем основной сценарий")
        draft_program = write_program_from_dossier(
            submission["questionnaire"],
            dossier,
            lambda percent, message, partial_program: save_generation_state(
                submission_id,
                status="running",
                stage="writer",
                percent=percent,
                message=message,
                program=finalize_generated_program(partial_program, submission["questionnaire"], dossier) if partial_program else None,
            ),
        )
        save_generation_state(submission_id, status="running", stage="polish", percent=78, message="Шлифуем тексты и DJ sheet")
        program = polish_program(
            submission["questionnaire"],
            dossier,
            draft_program,
            lambda percent, message, partial_program: save_generation_state(
                submission_id,
                status="running",
                stage="polish",
                percent=percent,
                message=message,
                program=finalize_generated_program(partial_program, submission["questionnaire"], dossier),
            ),
        )
        program = finalize_generated_program(program, submission["questionnaire"], dossier)
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"Ошибка генерации программы: {error}") from error

    submissions[index]["program"] = program
    submissions[index]["generation"] = {
        "status": "ready",
        "stage": "final_assembly",
        "percent": 100,
        "message": "Программа готова",
        "error": "",
        "updated_at": datetime.now().isoformat(),
    }
    save_submissions(submissions)
    return {"status": "success", "submissionId": submission_id, "program": program, "cached": False}


@app.get("/api/submissions/{submission_id}/export-docx")
def export_docx(submission_id: str) -> StreamingResponse:
    submissions, index = get_submission_or_404(submission_id)
    submission = submissions[index]

    program = submission.get("program")
    if not is_program_actual(program):
        try:
            start_generation_job(submission_id)
            raise HTTPException(status_code=409, detail="Программа еще формируется. Дождитесь завершения генерации и попробуйте скачать файл снова.")
        except HTTPException:
            raise
        except Exception as error:
            raise HTTPException(status_code=500, detail=f"Ошибка генерации перед экспортом: {error}") from error

    try:
        buffer = build_docx(submission, program)
        client_name = safe_filename(submission["questionnaire"].get("clientName", "client"))
        event_type = safe_filename(submission["questionnaire"].get("eventType", "event"))
        filename = f"{event_type}_{client_name}_{submission_id}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"Ошибка Word-выгрузки: {error}") from error
